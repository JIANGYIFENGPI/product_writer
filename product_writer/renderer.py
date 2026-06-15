from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from product_writer.bolding import (
    is_structure_line,
    leading_emphasis_end,
    split_by_terms,
    split_markdown_bold,
    strip_markdown_bold,
)
from product_writer.image_assets import ArticleImages, image_keywords

FONT_NAME = "宋体"
FONT_SIZE_PT = 12
IMAGE_WIDTH_IN = 6.0
DEFAULT_LONG_PREAMBLE_CHARS = 300
DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES = 2
DEFAULT_MIN_PARAGRAPHS_BEFORE_NEUTRAL = 2


def set_run_font(run, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def set_para(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = None


def fix_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = None


def add_formatted_text(paragraph, text: str, terms: list[str], bold_all: bool = False) -> None:
    if bold_all:
        run = paragraph.add_run(strip_markdown_bold(text).replace("*", ""))
        set_run_font(run, bold=True)
        return

    emphasis_end = leading_emphasis_end(strip_markdown_bold(text))
    plain_offset = 0
    for markdown_chunk, markdown_bold in split_markdown_bold(text):
        if markdown_bold:
            run = paragraph.add_run(markdown_chunk)
            set_run_font(run, bold=True)
            plain_offset += len(markdown_chunk)
            continue
        chunks: list[tuple[str, bool]] = []
        if emphasis_end > plain_offset:
            local_end = min(len(markdown_chunk), emphasis_end - plain_offset)
            if local_end > 0:
                chunks.append((markdown_chunk[:local_end], True))
            if local_end < len(markdown_chunk):
                chunks.extend(split_by_terms(markdown_chunk[local_end:], terms))
        else:
            chunks.extend(split_by_terms(markdown_chunk, terms))
        for chunk, term_bold in chunks:
            if not chunk:
                continue
            run = paragraph.add_run(chunk)
            set_run_font(run, bold=term_bold)
        plain_offset += len(markdown_chunk)


def add_plain_paragraph(doc: Document, text: str, terms: list[str], bold_all: bool = False) -> None:
    paragraph = doc.add_paragraph()
    set_para(paragraph)
    add_formatted_text(paragraph, text, terms, bold_all=bold_all)


def add_picture_paragraph(doc: Document, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    set_para(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(IMAGE_WIDTH_IN))


def ranking_number(text: str) -> int | None:
    stripped = strip_markdown_bold(text).strip().lstrip("#").strip()
    chinese_numbers = {
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    direct_prefixes = {
        "首推": 1,
        "次推": 2,
        "三推": 3,
        "第三推": 3,
    }
    for prefix, number in direct_prefixes.items():
        if stripped.startswith(prefix):
            return number
    patterns = [
        r"^TOP\s*(\d+)",
        r"^推荐([一二三四五六七八九十\d]+)",
        r"^第([一二三四五六七八九十\d]+)(?:名|款|位)",
        r"^([一二三])(?:、|[.．])",
    ]
    for pattern in patterns:
        match = re.match(pattern, stripped, flags=re.IGNORECASE)
        if not match:
            continue
        value = match.group(1)
        if value.isdigit():
            return int(value)
        return chinese_numbers.get(value)
    return None


def explicit_product_ranking_number(text: str) -> int | None:
    stripped = strip_markdown_bold(text).strip().lstrip("#").strip()
    if not re.match(
        r"^(?:TOP\s*\d+|推荐[一二三四五六七八九十\d]+|"
        r"第[一二三四五六七八九十\d]+(?:名|款|位))",
        stripped,
        flags=re.IGNORECASE,
    ):
        return None
    return ranking_number(stripped)


def promoted_product_rank(text: str, config: dict[str, Any]) -> int | None:
    stripped = strip_markdown_bold(text).strip().lstrip("#").strip()
    if len(stripped) > 120:
        return None
    matched: list[int] = []
    for item in config.get("promoted_products") or []:
        names = [item.get("brand"), item.get("product_name"), *(item.get("aliases") or [])]
        if any(name and str(name) in stripped for name in names):
            matched.append(int(item["rank"]))
    return matched[0] if len(set(matched)) == 1 else None


def find_detail_image_anchors(blocks: list[str], config: dict[str, Any]) -> dict[int, int]:
    """Return the first real detail block for each promoted product.

    Short consecutive ranking lines are summaries, not detail sections. A
    candidate only qualifies when it has substantive same-product prose after
    it. The fallback accepts a long product paragraph when the model omitted a
    dedicated heading.
    """
    plain_blocks = [strip_markdown_bold(block).strip().lstrip("#").strip() for block in blocks]
    candidates: dict[int, list[tuple[int, int]]] = {1: [], 2: [], 3: []}

    for index, plain in enumerate(plain_blocks):
        explicit_rank = explicit_product_ranking_number(plain)
        rank = promoted_product_rank(plain, config) or explicit_rank
        if rank not in {1, 2, 3}:
            continue
        explicit_heading = is_structure_line(plain) or explicit_rank == rank
        short_heading = len(plain) <= 100
        if not explicit_heading and not short_heading:
            continue

        detail_chars = 0
        for following in plain_blocks[index + 1:index + 6]:
            following_rank = promoted_product_rank(following, config)
            if following_rank is not None and following_rank != rank:
                break
            if explicit_product_ranking_number(following) not in {None, rank}:
                break
            if following_rank == rank and len(following) >= 80:
                detail_chars += len(following)
                continue
            if is_structure_line(following):
                continue
            detail_chars += len(following)
        if detail_chars >= 60:
            candidates[rank].append((0 if explicit_heading else 1, index))

    anchors: dict[int, int] = {}
    for rank in (1, 2, 3):
        if candidates[rank]:
            anchors[rank] = min(candidates[rank])[1]
            continue
        for index, plain in enumerate(plain_blocks):
            if promoted_product_rank(plain, config) == rank and len(plain) >= 180:
                anchors[rank] = index
                break
    return anchors


def neutral_image_anchor(
    blocks: list[str],
    first_detail_index: int | None,
    config: dict[str, Any],
    neutral_image: Path | None = None,
) -> tuple[str, int]:
    """Choose the neutral-image position before the first promoted product.

    Returns ("after", block_index) or ("before", block_index).
    """
    settings = config.get("illustrations") or {}
    configured_strategy = settings.get("neutral_placement_strategy")
    strategy = str(
        configured_strategy
        or (
            "fixed_body_count"
            if int(settings.get("cover_after_body_paragraphs") or 0) > 0
            else "section_boundary_with_spacing"
        )
    ).strip()
    if strategy == "section_boundary_with_spacing":
        if first_detail_index is None:
            return ("after", 0)
        minimum_chars = int(
            settings.get("neutral_body_paragraph_min_chars") or 60
        )
        min_before = int(
            settings.get("minimum_body_paragraphs_before_neutral") or 5
        )
        min_after = int(
            settings.get("minimum_body_paragraphs_between_images") or 3
        )
        body_indices = [
            index
            for index, block in enumerate(blocks[:first_detail_index])
            if len(strip_markdown_bold(block).strip()) >= minimum_chars
            and not is_structure_line(strip_markdown_bold(block))
        ]
        boundaries: list[tuple[int, int, int]] = []
        for position, index in enumerate(body_indices):
            paragraphs_before = position + 1
            paragraphs_after = len(body_indices) - paragraphs_before
            if paragraphs_before < min_before:
                continue
            next_index = index + 1
            while next_index < first_detail_index and not blocks[next_index].strip():
                next_index += 1
            next_plain = (
                strip_markdown_bold(blocks[next_index]).strip()
                if next_index < first_detail_index
                else ""
            )
            following_plain = (
                strip_markdown_bold(blocks[next_index + 1]).strip()
                if next_index + 1 < first_detail_index
                else ""
            )
            contextual_boundary = bool(
                2 <= len(next_plain) <= 36
                and not any(mark in next_plain for mark in "。！？!?；;")
                and len(following_plain) >= minimum_chars
                and explicit_product_ranking_number(next_plain) is None
            )
            if (
                next_index < first_detail_index
                and (
                    is_structure_line(next_plain)
                    or _configured_structure_line(next_plain, config)
                    or contextual_boundary
                )
            ):
                boundaries.append((position, index, paragraphs_after))
        eligible = [
            (position, index)
            for position, index, paragraphs_after in boundaries
            if paragraphs_after >= min_after
        ]
        if eligible:
            target = max(min_before - 1, int((len(body_indices) - 1) * 0.6))
            _, selected = min(
                eligible,
                key=lambda item: (abs(item[0] - target), -item[0]),
            )
            return ("after", selected)
        if boundaries:
            _, selected, _ = max(boundaries, key=lambda item: item[2])
            return ("after", selected)
        return ("before", first_detail_index)
    if strategy == "balanced_body_with_spacing":
        if first_detail_index is None:
            return ("after", 0)
        minimum_chars = int(
            settings.get("neutral_body_paragraph_min_chars") or 80
        )
        min_before = int(
            settings.get("minimum_body_paragraphs_before_neutral") or 2
        )
        min_after = int(
            settings.get("minimum_body_paragraphs_between_images")
            or DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES
        )
        body_indices = [
            index
            for index, block in enumerate(blocks[:first_detail_index])
            if len(strip_markdown_bold(block).strip()) >= minimum_chars
            and not is_structure_line(strip_markdown_bold(block))
        ]
        if len(body_indices) >= min_before + min_after:
            latest_position = len(body_indices) - min_after - 1
            target_position = min(
                latest_position,
                max(min_before - 1, (len(body_indices) - 1) // 2),
            )
            return ("after", body_indices[target_position])
        if len(body_indices) >= min_before:
            return ("after", body_indices[min_before - 1])
        if body_indices:
            return ("after", body_indices[-1])
        return ("after", 0)
    if strategy == "first_body_with_spacing":
        if first_detail_index is None:
            return ("after", 0)
        minimum_chars = int(
            settings.get("neutral_body_paragraph_min_chars") or 80
        )
        long_preamble_chars = int(
            settings.get("neutral_after_first_paragraph_when_preamble_chars")
            or DEFAULT_LONG_PREAMBLE_CHARS
        )
        min_paragraphs = int(
            settings.get("minimum_body_paragraphs_between_images")
            or DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES
        )
        body_indices = [
            index
            for index, block in enumerate(blocks[:first_detail_index])
            if len(strip_markdown_bold(block).strip()) >= minimum_chars
            and not is_structure_line(strip_markdown_bold(block))
        ]
        preamble_chars = sum(
            len(strip_markdown_bold(blocks[index]).strip())
            for index in body_indices
        )
        if (
            body_indices
            and len(body_indices) >= min_paragraphs + 1
            and preamble_chars >= long_preamble_chars
        ):
            return ("after", body_indices[0])
        if body_indices:
            return ("before", body_indices[0])
        return ("after", 0)

    configured_target = int(settings.get("cover_after_body_paragraphs") or 0)
    if configured_target > 0:
        search_end = first_detail_index if first_detail_index is not None else len(blocks)
        minimum_chars = max(
            40,
            int(settings.get("cover_body_paragraph_min_chars") or 80),
        )
        body_indices: list[int] = []
        for index, block in enumerate(blocks[:search_end]):
            plain = strip_markdown_bold(block).strip()
            if len(plain) < minimum_chars:
                continue
            if is_structure_line(plain):
                continue
            if explicit_product_ranking_number(plain) is not None:
                continue
            body_indices.append(index)
            if len(body_indices) >= configured_target:
                return ("after", index)

        if body_indices:
            return ("after", body_indices[-1])
        for index, block in enumerate(blocks[:search_end]):
            plain = strip_markdown_bold(block).strip()
            if plain and not is_structure_line(plain):
                return ("after", index)
        return ("after", 0)

    # Optional legacy semantic placement strategy.
    if first_detail_index is None:
        return ("after", 0)
    long_preamble_chars = int(
        settings.get("neutral_after_first_paragraph_when_preamble_chars")
        or DEFAULT_LONG_PREAMBLE_CHARS
    )
    min_paragraphs = int(
        settings.get("minimum_body_paragraphs_between_images")
        or DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES
    )
    min_before = int(
        settings.get("minimum_body_paragraphs_before_neutral")
        or DEFAULT_MIN_PARAGRAPHS_BEFORE_NEUTRAL
    )
    body_indices = [
        index
        for index, block in enumerate(blocks[:first_detail_index])
        if len(strip_markdown_bold(block).strip()) >= 40
        and not is_structure_line(strip_markdown_bold(block))
    ]
    preamble_chars = sum(len(strip_markdown_bold(blocks[index])) for index in body_indices)

    if neutral_image and settings.get("semantic_neutral_anchor", True):
        keywords = image_keywords(
            neutral_image,
            settings.get("neutral_image_keywords") or {},
        )
        scored: list[tuple[int, int]] = []
        for position, index in enumerate(body_indices):
            paragraphs_before = position + 1
            paragraphs_after = len(body_indices) - position - 1
            if paragraphs_before < min_before or paragraphs_after < min_paragraphs:
                continue
            plain = strip_markdown_bold(blocks[index]).lower()
            score = sum(
                min(plain.count(keyword.lower()), 3) * max(1, min(len(keyword), 6))
                for keyword in keywords
                if keyword
            )
            if score > 0:
                scored.append((score, index))
        if scored:
            best_score = max(score for score, _ in scored)
            best_indices = [index for score, index in scored if score == best_score]
            return ("after", best_indices[0])

    if (
        body_indices
        and len(body_indices) >= min_before + min_paragraphs
        and preamble_chars >= long_preamble_chars
    ):
        target_position = min(
            max(min_before - 1, (len(body_indices) - 1) // 2),
            len(body_indices) - min_paragraphs - 1,
        )
        return ("after", body_indices[target_position])
    if body_indices:
        return ("after", body_indices[-1])
    return ("after", 0)


def _normalize_title_for_compare(text: str) -> str:
    plain = strip_markdown_bold(text).strip()
    plain = plain.lstrip("#").strip()
    plain = plain.strip("《》「」『』“”\"'")
    plain = re.sub(r"\s+", "", plain)
    plain = plain.replace("：", ":")
    return plain


def _strip_duplicate_title(body: str, title: str) -> str:
    """移除正文开头与标题重复的行。

    模型经常在正文第一行重复输出标题（可能带 markdown 加粗、# 前缀等变体）。
    检查正文前几行，找到并移除与标题相同的行。比较时忽略空格、冒号差异、
    markdown 标记和外层书名号。
    """
    plain_title = strip_markdown_bold(title).strip()
    normalized_title = _normalize_title_for_compare(title)
    lines = body.split("\n")

    # 在前 5 个非空行中查找标题重复
    removed = 0
    keep_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if removed < 3 and stripped:
            plain = strip_markdown_bold(stripped).strip()
            normalized_plain = _normalize_title_for_compare(stripped)
            if normalized_plain == normalized_title:
                removed += 1
                continue
            if normalized_plain.startswith(normalized_title) and len(normalized_plain) <= len(normalized_title) + 5:
                suffix = normalized_plain[len(normalized_title):].lstrip(":")
                if not suffix or suffix == normalized_title:
                    removed += 1
                    continue
            if plain == plain_title:
                removed += 1
                continue
            if plain.startswith(plain_title) and len(plain) <= len(plain_title) + 5:
                suffix = plain[len(plain_title):].lstrip("：: \t")
                if not suffix or suffix == plain_title:
                    removed += 1
                    continue
        keep_lines.append(line)

    if removed:
        return "\n".join(keep_lines).strip()

    # fallback: 精确字符串前缀匹配
    stripped_body = body.strip()
    if stripped_body.startswith(title):
        return stripped_body[len(title):].lstrip(" \n\r\t:：")

    return body


def _configured_structure_line(text: str, config: dict[str, Any]) -> bool:
    plain = strip_markdown_bold(text).strip().strip("【】[]")
    headings = (config.get("article_structure") or {}).get("headings") or []
    return any(
        plain == strip_markdown_bold(str(heading)).strip().strip("【】[]")
        for heading in headings
        if str(heading).strip()
    )


def render_docx(
    title: str,
    text: str,
    output_path: Path,
    config: dict[str, Any],
    terms: list[str],
    image_paths: list[Path] | ArticleImages | None = None,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    fix_normal_style(doc)

    add_plain_paragraph(doc, title, terms=[], bold_all=True)

    body = text.strip()
    body = _strip_duplicate_title(body, title)

    active_terms = terms if config["features"].get("bold_terms", True) else []
    if isinstance(image_paths, ArticleImages):
        article_images = image_paths
    else:
        legacy = image_paths or []
        article_images = ArticleImages(
            neutral=legacy[0] if legacy else None,
            by_rank={
                rank: legacy[rank]
                for rank in (1, 2, 3)
                if rank < len(legacy)
            },
        )
    blocks = [part.strip() for part in body.split("\n") if part.strip()]
    detail_anchors = find_detail_image_anchors(blocks, config)
    first_detail_index = min(detail_anchors.values()) if detail_anchors else None
    neutral_position = neutral_image_anchor(
        blocks,
        first_detail_index,
        config,
        article_images.neutral,
    )

    in_qa_section = False
    inserted_extra_ranks: set[int] = set()
    for index, block in enumerate(blocks):
        if neutral_position == ("before", index) and article_images.neutral:
            add_picture_paragraph(doc, article_images.neutral)
        plain_block = strip_markdown_bold(block)
        configured_heading = _configured_structure_line(plain_block, config)
        next_plain = (
            strip_markdown_bold(blocks[index + 1]).strip()
            if index + 1 < len(blocks)
            else ""
        )
        contextual_heading = bool(
            2 <= len(plain_block.strip()) <= 36
            and not any(mark in plain_block for mark in "。！？!?；;")
            and len(next_plain) >= 60
            and not explicit_product_ranking_number(plain_block)
        )
        structure_heading = (
            is_structure_line(plain_block)
            or configured_heading
            or contextual_heading
        )
        if structure_heading:
            if any(marker in plain_block for marker in ("常见问题", "答疑", "问答")):
                in_qa_section = True
            elif any(marker in plain_block for marker in ("总结", "结语", "写在最后")):
                in_qa_section = False
        bold_all = bool(
            config["features"].get("bold_structure", True)
            and structure_heading
        )
        add_plain_paragraph(doc, block, active_terms, bold_all=bold_all)
        if neutral_position == ("after", index) and article_images.neutral:
            add_picture_paragraph(doc, article_images.neutral)
        for rank in (1, 2, 3):
            if detail_anchors.get(rank) == index and rank in article_images.by_rank:
                add_picture_paragraph(doc, article_images.by_rank[rank])
        for rank, extra_images in article_images.extra_by_rank.items():
            if not extra_images or rank in inserted_extra_ranks:
                continue
            anchor = detail_anchors.get(rank)
            if anchor is None or index <= anchor:
                continue
            plain = strip_markdown_bold(block).strip()
            if is_structure_line(plain) or len(plain) < 80:
                continue
            for image_path in extra_images:
                add_picture_paragraph(doc, image_path)
            inserted_extra_ranks.add(rank)

    for paragraph in doc.paragraphs:
        set_para(paragraph)
        for run in paragraph.runs:
            set_run_font(run, bold=bool(run.bold))

    doc.save(output_path)
