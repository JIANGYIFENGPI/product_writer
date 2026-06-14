from __future__ import annotations

from collections import Counter
from difflib import SequenceMatcher
import json
import re
from statistics import mean, pstdev
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from product_writer.bolding import is_structure_line
from product_writer.cleaner import contains_boilerplate
from product_writer.image_assets import promoted_image_slots
from product_writer.text_metrics import article_char_count

END_PUNCTUATION = set("。！？!?」』）)]”’")
AI_SELF_PROBES = [
    "作为一名",
    "作为专业",
    "作为AI",
    "作为一个AI",
    "我是AI",
    "我将为您",
    "我将严格",
    "我会按照",
    "根据您的要求",
    "根据你的要求",
    "按您的要求",
    "按你的要求",
    "以下是",
    "下面是",
    "下面的内容",
    "这篇内容",
    "下文将",
    "接下来将",
    "目的是让",
    "本篇将",
    "本篇整理",
    "本篇内容",
    "本文将",
    "文章将",
    "这篇文章",
    "希望本文",
    "希望这篇",
    "抱歉",
    "无法提供",
]
PROMPT_ECHO_PROBES = [
    "最高优先级",
    "硬规则",
    "固定要求",
    "固定TOP",
    "固定 TOP",
    "固定的TOP",
    "固定的 TOP",
    "TOP1-TOP10排行榜",
    "TOP1-TOP10 排行榜",
    "只使用您提供",
    "只使用你提供",
    "用户资料中",
    "按照固定榜单",
    "遵循您的所有指令",
    "遵循你的所有指令",
    "品牌白名单",
    "白名单品牌",
]
TEMPLATE_TONE_PROBES = [
    "在当今快节奏生活中",
    "随着人们健康意识提升",
    "随着生活水平的提高",
    "在现代社会",
    "在当前市场环境下",
    "本文将从多个维度",
    "通过本文你将了解到",
    "通过本文您将了解到",
    "综上所述",
    "总而言之",
    "总的来说",
    "由此可见",
    "不难发现",
    "值得一提的是",
    "需要注意的是",
    "值得注意的是",
    "值得关注的是",
    "更重要的是",
    "更关键的是",
    "需要反复强调的是",
    "信息迷雾",
    "建立一种判断框架",
    "建立一套判断框架",
    "本文围绕",
    "本文从",
    "消费者需要",
    "消费者应当",
    "购买前应",
    "购买前需要",
    "从配料表来看",
    "从产品名称来看",
    "从使用场景来看",
    "这意味着",
    "毋庸置疑",
]
MECHANICAL_CONNECTORS = [
    "首先",
    "其次",
    "再次",
    "最后",
    "此外",
    "同时",
    "因此",
    "所以",
    "一方面",
    "另一方面",
]
VAGUE_ATTRIBUTION_PATTERNS = [
    r"(?:有|据)研究(?:表明|显示|发现)",
    r"数据显示",
    r"调查显示",
    r"专家(?:认为|表示|指出|建议)",
    r"业内人士(?:认为|表示|指出)",
    r"权威机构(?:认为|表示|指出|发布)",
    r"相关资料(?:表明|显示)",
]
SCOPE_OVERSTATEMENT_PROBES = [
    "覆盖了市场上各大主流需求",
    "覆盖市场上各大主流需求",
    "覆盖整个市场",
    "囊括市场所有品牌",
    "覆盖全部品牌",
]
BRAND_SECTION_STOP_PREFIXES = (
    "读完",
    "接下来",
    "按需而选",
    "避坑指南",
    "不同人群",
    "人群专属",
    "选购建议",
    "行业乱象",
    "常见问题",
    "常见问题快问快答",
    "常见疑问",
    "核心答疑",
    "消费倡议",
    "总结建议",
)
EXPECTED_FONT_NAME = "宋体"
EXPECTED_FONT_SIZE_PT = 12
IMAGE_PLACEHOLDER_PROBES = [
    "此处插图",
    "此处配图",
    "图片位置",
    "插入图片",
    "配图建议",
    "[图片",
    "【图片",
]
EDITORIAL_DISCLAIMER_PROBES = [
    "包装名称",
    "实物标签",
    "核对标签",
    "从名称来看",
    "名称本身",
    "以实物为准",
    "信息还需确认",
    "只能依据产品名称",
    "不能仅凭名称",
    "产品名称直接",
    "产品名称突出",
    "名称直接突出",
    "从产品名称",
    "产品路线",
    "不构成整个",
    "不是必选项",
    "形成两条不同",
    "选择哪条路线",
    "谁更全",
    "代表不同思路",
    "代表了两种",
    "取决于自己日常饮食",
    "营养相对不足",
    "补足饮食短板",
    "以上这些判断方法",
    "没有绑定具体品牌",
    "套到具体产品上",
    "逐一推敲",
]
RANKING_DISCLAIMER_PROBES = [
    "序号不代表市场排名",
    "序号不代表排名",
    "不代表市场排名或优劣",
    "不代表产品优劣",
    "不是商业排名",
    "仅为陈列顺序",
    "仅作选购信息参考",
]
PRODUCT_EDITORIAL_PROBES = [
    "资料中标出",
    "工艺资料",
    "吸收率资料",
    "资料说明",
    "已有资料",
    "销量资料",
    "它的资料",
    "资料重点",
    "资料没有",
    "信息完整度",
    "标签语言",
    "自行补充没有提供的结论",
    "不能因为共享一个成分",
    "视为同一种体验",
]
UNSUPPORTED_INFERENCE_PROBES = [
    "进入血液循环更容易",
    "胃排空速度比片剂快",
    "降低肠胃的工作负荷",
    "对胃肠道刺激性更低",
    "吸收曲线更平缓",
    "有助于维持黏膜湿润",
    "不会成为发胖的来源",
    "不存在依赖性问题",
    "回到饮前状态",
    "减少买到临期",
    "降低买到临期",
    "减少未被吸收就排出",
    "等量摄入下实际被利用",
    "晚间喝可以和身体夜间",
    "午后喝则能弥补",
]
DATA_CAUSAL_PROBES = [
    "复购率更能体现饮用体验",
    "复购率说明实际体验",
    "高复购率至少说明",
    "销量反映市场接受度",
    "意味着品牌在品控",
    "至少说明品牌在品控",
    "说明市场接受度和持续选择意愿",
    "认证跨度覆盖国内与国际",
]


def contains_markdown_table(text: str) -> bool:
    """检测文本中是否包含 markdown 表格（连续的 | 分隔行）。"""
    lines = text.split("\n")
    pipe_lines = 0
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and "|" in stripped[1:]:
            pipe_lines += 1
            if pipe_lines >= 3:
                return True
        else:
            pipe_lines = 0
    return False


def output_delivery_warnings(text: str) -> list[str]:
    warnings: list[str] = []
    image_probes = [probe for probe in IMAGE_PLACEHOLDER_PROBES if probe in text]
    if image_probes:
        warnings.append("正文包含图片占位或配图说明：" + "、".join(image_probes))
    if re.search(r"(?m)^[ \t\u3000]+", text):
        warnings.append("正文存在行首空格或缩进")
    if re.search(r"(?m)^\s*#{1,6}\s+", text):
        warnings.append("正文包含 Markdown 标题符号")
    if re.search(r"(?m)^\s*[-+*]\s+", text):
        warnings.append("正文包含项目符号列表")
    return warnings


def humanizer_warnings(text: str, config: dict[str, Any] | None = None) -> list[str]:
    warnings: list[str] = []
    disclaimer_probes = list(EDITORIAL_DISCLAIMER_PROBES)
    layout = (config or {}).get("recommendation_layout") or {}
    if layout.get("reject_ranking_disclaimer", False):
        disclaimer_probes.extend(RANKING_DISCLAIMER_PROBES)
    disclaimer_hits = [probe for probe in disclaimer_probes if probe in text]
    if disclaimer_hits:
        warnings.append(
            "包含资料不足或标签核对式AI套话："
            + "、".join(disclaimer_hits[:8])
        )

    self_found = [probe for probe in AI_SELF_PROBES if probe in text]
    if self_found:
        warnings.append("包含AI自我说明或客套口吻：" + "、".join(self_found[:8]))

    prompt_found = [probe for probe in PROMPT_ECHO_PROBES if probe in text]
    if prompt_found:
        warnings.append("包含提示词或规则复述痕迹：" + "、".join(prompt_found[:8]))

    template_counts = {
        probe: text.count(probe)
        for probe in TEMPLATE_TONE_PROBES
        if probe in text
    }
    template_total = sum(template_counts.values())
    repeated_templates = [
        f"{probe}×{count}"
        for probe, count in template_counts.items()
        if count >= 3
    ]
    if len(template_counts) >= 3 or template_total >= 5 or repeated_templates:
        detail = repeated_templates or list(template_counts)
        warnings.append("包含模板化AI腔表达：" + "、".join(detail[:8]))

    connector_hits = {
        connector: len(re.findall(rf"(?m)(?:^|[。！？；;]\s*){re.escape(connector)}[，,]", text))
        for connector in MECHANICAL_CONNECTORS
    }
    repeated = [f"{connector}×{count}" for connector, count in connector_hits.items() if count >= 4]
    total = sum(connector_hits.values())
    if repeated or total >= 12:
        detail = "、".join(repeated[:6]) if repeated else f"连接词总数×{total}"
        warnings.append("连接词使用过于机械：" + detail)

    vague_hits = []
    for pattern in VAGUE_ATTRIBUTION_PATTERNS:
        vague_hits.extend(re.findall(pattern, text))
    if len(vague_hits) >= 2:
        warnings.append("包含多处无明确来源的模糊归因：" + "、".join(vague_hits[:6]))

    paragraphs = [
        article_char_count(line)
        for line in text.splitlines()
        if article_char_count(line) >= 40
    ]
    if len(paragraphs) >= 8:
        average = mean(paragraphs)
        variation = pstdev(paragraphs) / average if average else 0
        if variation < 0.12:
            warnings.append("正文段落长度过度整齐，疑似模板化生成")

    body_lines = [
        re.sub(r"^[“\"'‘]+", "", line.strip())
        for line in text.splitlines()
        if article_char_count(line) >= 45
    ]
    opening_counts: dict[str, int] = {}
    for line in body_lines:
        opening = re.sub(r"^[（(]?[一二三四五六七八九十\d]+[）)、.．\s]*", "", line)
        opening = re.sub(r"^(?:对于|关于)", "", opening)
        prefix = opening[:5]
        if len(prefix) == 5:
            opening_counts[prefix] = opening_counts.get(prefix, 0) + 1
    opening_repeat_limit = 7
    repeated_openings = [
        f"{prefix}…×{count}"
        for prefix, count in sorted(
            opening_counts.items(),
            key=lambda item: (-item[1], item[0]),
        )
        if count >= opening_repeat_limit
    ]
    if repeated_openings:
        warnings.append("多段使用相同开头，行文过于程式化：" + "、".join(repeated_openings[:6]))

    comparable_lines = [
        re.sub(r"[，。！？；：、“”‘’（）()\s]", "", line.strip())
        for line in text.splitlines()
        if article_char_count(line) >= 55 and not is_structure_line(line.strip())
    ]
    near_duplicates: list[str] = []
    for index, current in enumerate(comparable_lines):
        for later in comparable_lines[index + 1 : index + 4]:
            shorter = min(len(current), len(later))
            longer = max(len(current), len(later))
            if shorter / longer < 0.65:
                continue
            similarity = SequenceMatcher(None, current, later).ratio()
            if similarity >= 0.56:
                near_duplicates.append(f"相邻段落相似度{similarity:.0%}")
                break
    if near_duplicates:
        warnings.append(
            "存在近义重复段落，应合并而不是换词复述："
            + "、".join(near_duplicates[:4])
        )

    strong_ai_patterns = {
        "不是……而是……": r"不是[^。！？]{0,35}而是",
        "本质上": r"本质上",
        "底层逻辑": r"底层逻辑",
        "真正关键": r"真正(?:关键|重要|值得|能)",
        "这意味着": r"这意味着",
        "最后一公里": r"最后一公里",
        "试金石": r"试金石",
    }
    strong_ai_hits = {
        label: len(re.findall(pattern, text))
        for label, pattern in strong_ai_patterns.items()
    }
    strong_ai_hits = {label: count for label, count in strong_ai_hits.items() if count}
    if sum(strong_ai_hits.values()) >= 5:
        warnings.append(
            "抽象判断和转折套话过多："
            + "、".join(f"{label}×{count}" for label, count in strong_ai_hits.items())
        )

    checklist_probes = [
        "核对配料表",
        "查看配料表",
        "核对实物标签",
        "以实物标签为准",
        "生产许可证",
        "建议食用量",
        "储存方式",
    ]
    checklist_counts = {
        probe: text.count(probe)
        for probe in checklist_probes
        if text.count(probe) >= 5
    }
    if len(checklist_counts) >= 2:
        warnings.append(
            "标签核验提醒重复过多，推荐段落疑似套用同一模板："
            + "、".join(f"{probe}×{count}" for probe, count in checklist_counts.items())
        )

    recommendation_starts = [
        re.sub(r"^[“\"'‘]+", "", line.strip())[:9]
        for line in text.splitlines()
        if article_char_count(line) >= 55
        and any(
            probe in line[:24]
            for probe in (
                "选择这款",
                "选择该产品",
                "关注这款",
                "查看这款",
                "对于这款",
                "购买这款",
                "这款产品",
                "这类产品",
            )
        )
    ]
    recommendation_counts = Counter(recommendation_starts)
    repeated_recommendations = [
        f"{opening}…×{count}"
        for opening, count in recommendation_counts.items()
        if count >= 3
    ]
    if repeated_recommendations:
        warnings.append(
            "多款产品介绍使用相同套话开头："
            + "、".join(repeated_recommendations[:6])
        )

    product_editorial_hits: list[str] = []
    in_product_section = False
    seen_product_body = False
    for line in text.splitlines():
        stripped = line.strip().lstrip("#* ")
        if _is_recommendation_heading(stripped):
            in_product_section = True
            seen_product_body = False
            continue
        if in_product_section:
            if _ends_recommendation_section(stripped) or (
                seen_product_body and is_structure_line(stripped)
            ):
                in_product_section = False
            elif stripped and not is_structure_line(stripped):
                seen_product_body = True
        if not in_product_section:
            continue
        product_editorial_hits.extend(
            probe for probe in PRODUCT_EDITORIAL_PROBES if probe in stripped
        )
    if product_editorial_hits:
        warnings.append(
            "产品介绍包含内部资料或编辑规则口吻："
            + "、".join(dict.fromkeys(product_editorial_hits))
        )
    unsupported_hits = [
        probe for probe in UNSUPPORTED_INFERENCE_PROBES if probe in text
    ]
    if unsupported_hits:
        warnings.append(
            "包含无资料支撑的医学或效果推断："
            + "、".join(unsupported_hits[:8])
        )
    causal_hits = [probe for probe in DATA_CAUSAL_PROBES if probe in text]
    if causal_hits:
        warnings.append(
            "将销量、复购或认证数据强行解释为体验或效果："
            + "、".join(causal_hits[:8])
        )
    return warnings


def ai_trace_warnings(
    text: str,
    config: dict[str, Any] | None = None,
) -> list[str]:
    return humanizer_warnings(text, config)


def complete_structure_warnings(text: str, config: dict[str, Any]) -> list[str]:
    settings = config.get("article_structure") or {}
    if not settings.get("enabled", False):
        return []
    groups = settings.get("required_heading_groups") or []
    heading_lines = [
        line.strip()
        for line in text.splitlines()
        if 2 <= article_char_count(line.strip()) <= 80
        and is_structure_line(line.strip())
    ]
    minimum_heading_count = int(settings.get("minimum_heading_count") or 0)
    warnings: list[str] = []
    if minimum_heading_count and len(heading_lines) < minimum_heading_count:
        warnings.append(
            f"独立结构标题不足：{len(heading_lines)} < {minimum_heading_count}"
        )
    missing = []
    for group in groups:
        keywords = [str(item).strip() for item in group if str(item).strip()]
        if keywords and not any(
            any(keyword.lower() in line.lower() for keyword in keywords)
            for line in heading_lines
        ):
            missing.append("/".join(keywords))
    if missing:
        warnings.append("缺少独立结构标题：" + "、".join(missing))
    return warnings


def paragraph_structure_warnings(text: str) -> list[str]:
    warnings: list[str] = []
    dense_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        sentence_count = len(re.findall(r"[。！？!?]", stripped))
        if article_char_count(stripped) >= 260 and sentence_count >= 3:
            dense_lines.append(stripped[:45])
    if dense_lines:
        warnings.append(
            f"正文存在{len(dense_lines)}个包含多个观点的超长段落，需按句意拆段："
            + "；".join(dense_lines[:3])
        )
    point_dense = []
    for line in text.splitlines():
        stripped = line.strip()
        enum_count = len(
            re.findall(
                r"(?:其[一二三四五六七八九十]是?|[一二三四五六七八九十]是)",
                stripped,
            )
        )
        if article_char_count(stripped) >= 160 and (
            stripped.count("；") >= 2 or enum_count >= 2
        ):
            point_dense.append(stripped[:45])
    if point_dense:
        warnings.append(
            f"正文存在{len(point_dense)}个合并书写的要点段落，需将分号或序号要点逐项另起段："
            + "；".join(point_dense[:3])
        )

    scope_hits = [probe for probe in SCOPE_OVERSTATEMENT_PROBES if probe in text]
    if scope_hits:
        warnings.append("十款榜单存在范围夸大表述：" + "、".join(scope_hits))
    return warnings


def repeated_numeric_claim_warnings(text: str) -> list[str]:
    repeated: list[str] = []
    claim_patterns = (
        r"吸收率做到95%以上",
        r"胶原蛋白肽吸收率为?95%以上",
        r"产品标准代号GB7101",
        r"食品生产许可证编号SC\d+",
        r"30日复购率(?:达到|超过|为)?85%",
        r"复购率(?:达到|超过|为)?90%以上",
    )
    for pattern in claim_patterns:
        matches = re.findall(pattern, text)
        if len(matches) >= 3:
            repeated.append(f"{matches[0]}×{len(matches)}")
    if repeated:
        return ["同一数字事实跨章节重复：" + "、".join(repeated)]
    return []


def visible_char_count(text: str) -> int:
    return article_char_count(text)


def conflicting_title_warnings(text: str, title: str) -> list[str]:
    lines = [line.strip().lstrip("#* ") for line in text.splitlines() if line.strip()]
    title_key = re.sub(r"\s+", "", title)
    conflicts: list[str] = []
    for line in lines[:5]:
        line_key = re.sub(r"\s+", "", line)
        if line_key == title_key:
            continue
        title_similarity = SequenceMatcher(None, title_key, line_key).ratio()
        if (
            12 <= len(line) <= 80
            and title_similarity >= 0.62
            and any(keyword in line for keyword in ("胶原蛋白", "驼奶", "产品", "品牌"))
        ):
            conflicts.append(line)
            continue
        if (
            18 <= len(line) <= 80
            and re.search(r"20\d{2}", line)
            and any(mark in line for mark in "！？?!")
        ):
            conflicts.append(line)
    if not conflicts:
        return []
    return ["正文开头出现与指定标题冲突的旧标题或副标题：" + "；".join(conflicts)]


def looks_truncated(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if stripped[-1] in END_PUNCTUATION:
        return False
    if stripped[-1] in "，,；;：:":
        return True
    return len(stripped.splitlines()[-1]) < 12


def inspect_docx(path: Path) -> tuple[bool, str | None]:
    try:
        Document(path)
        return True, None
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)


def _points(value: Any) -> float | None:
    if value is None:
        return None
    return round(float(value.pt), 2)


def docx_format_warnings(path: Path) -> list[str]:
    document = Document(path)
    warnings: list[str] = []
    empty_paragraphs: list[int] = []
    non_normal: list[int] = []
    bad_alignment: list[int] = []
    bad_spacing: list[int] = []
    bad_indent: list[int] = []
    bad_runs: list[str] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs, 1):
        has_picture = bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))
        if not paragraph.text.strip() and not has_picture:
            empty_paragraphs.append(paragraph_index)
        if paragraph.style.name != "Normal":
            non_normal.append(paragraph_index)

        pf = paragraph.paragraph_format
        if pf.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
            bad_alignment.append(paragraph_index)
        if (
            _points(pf.space_before) not in {None, 0.0}
            or _points(pf.space_after) not in {None, 0.0}
            or pf.line_spacing_rule != WD_LINE_SPACING.SINGLE
        ):
            bad_spacing.append(paragraph_index)
        if _points(pf.first_line_indent) not in {None, 0.0}:
            bad_indent.append(paragraph_index)

        for run_index, run in enumerate(paragraph.runs, 1):
            if not run.text:
                continue
            r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            font_names = {
                run.font.name,
                r_fonts.get(qn("w:ascii")),
                r_fonts.get(qn("w:hAnsi")),
                r_fonts.get(qn("w:eastAsia")),
            }
            if font_names != {EXPECTED_FONT_NAME} or _points(run.font.size) != EXPECTED_FONT_SIZE_PT:
                bad_runs.append(f"第{paragraph_index}段第{run_index}个文本块")

    if empty_paragraphs:
        warnings.append("Word 包含额外空段落：" + "、".join(map(str, empty_paragraphs[:10])))
    if non_normal:
        warnings.append("Word 包含非普通段落样式：" + "、".join(map(str, non_normal[:10])))
    if bad_alignment:
        warnings.append("Word 包含非左对齐段落：" + "、".join(map(str, bad_alignment[:10])))
    if bad_spacing:
        warnings.append("Word 段距或行距不符合规则：" + "、".join(map(str, bad_spacing[:10])))
    if bad_indent:
        warnings.append("Word 包含首行缩进：" + "、".join(map(str, bad_indent[:10])))
    if bad_runs:
        warnings.append(
            "Word 文本未统一设置宋体 12pt 及四类字体映射：" + "、".join(bad_runs[:10])
        )
    return warnings


def docx_article_char_count(path: Path) -> int:
    document = Document(path)
    return article_char_count("\n".join(paragraph.text for paragraph in document.paragraphs))


def docx_body_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs[1:] if paragraphs else []).strip()


def docx_illustration_warnings(path: Path, config: dict[str, Any]) -> list[str]:
    if not config.get("features", {}).get("images", False):
        return []
    document = Document(path)
    paragraphs = list(document.paragraphs)
    image_indices = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._element.xpath(".//w:drawing | .//w:pict")
    ]
    settings = config.get("illustrations") or {}
    required_count = int(
        settings.get("required_count")
        or (1 + len(promoted_image_slots(config)))
    )
    warnings: list[str] = []
    if len(image_indices) != required_count:
        return [f"插图数量异常：{len(image_indices)} != {required_count}"]
    if not settings.get("validate_positions", True):
        return []

    promoted = {
        int(item.get("rank") or 0): item
        for item in config.get("promoted_products") or []
        if int(item.get("rank") or 0) in promoted_image_slots(config)
    }
    top_indices: dict[int, int] = {}
    chinese_ranks = {
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        match = re.match(r"^TOP\s*(\d+)\s*[：:]", text, flags=re.IGNORECASE)
        if match:
            top_indices[int(match.group(1))] = index
            continue
        match = re.match(r"^推荐([一二三四五六七八九十\d]+)\s*[：:]", text)
        if match:
            value = match.group(1)
            rank = int(value) if value.isdigit() else chinese_ranks.get(value)
            if rank is not None:
                top_indices[rank] = index
                continue
        if len(text) > 120:
            continue
        for rank, item in promoted.items():
            names = [
                item.get("brand"),
                item.get("product_name"),
                *(item.get("aliases") or []),
            ]
            if any(name and str(name) in text for name in names):
                top_indices.setdefault(rank, index)

    promoted_ranks = sorted(promoted_image_slots(config))
    if promoted_ranks:
        first_rank = promoted_ranks[0]
        if first_rank not in top_indices or image_indices[0] >= top_indices[first_rank]:
            warnings.append("通用图必须位于首个主推产品详情之前")
        minimum_before = int(settings.get("minimum_body_paragraphs_before_neutral") or 2)
        body_before = sum(
            1
            for paragraph in paragraphs[1:image_indices[0]]
            if article_char_count(paragraph.text) >= 40
        )
        if body_before < minimum_before:
            warnings.append(
                f"标题与通用图之间正文段落不足：{body_before} < {minimum_before}"
            )
        for image_slot, rank in enumerate(promoted_ranks, 1):
            if rank not in top_indices or image_indices[image_slot] != top_indices[rank] + 1:
                warnings.append(f"第{rank}名产品图必须紧跟对应产品标题")
        minimum = int(settings.get("minimum_body_paragraphs_between_images") or 2)
        body_between = sum(
            1
            for paragraph in paragraphs[image_indices[0] + 1:image_indices[1]]
            if article_char_count(paragraph.text) >= 40
        )
        if body_between < minimum:
            warnings.append(
                f"通用图与首个主推产品图之间正文段落不足：{body_between} < {minimum}"
            )
    return warnings


def allowed_brand_names(config: dict[str, Any]) -> set[str]:
    names: set[str] = set()
    for item in config.get("brand_whitelist") or []:
        for key in ("brand", "product_name"):
            value = item.get(key)
            if value:
                names.add(str(value).strip())
    return {name for name in names if name}


def brand_mentions_before_recommendations_warnings(
    text: str,
    config: dict[str, Any],
) -> list[str]:
    lines = text.splitlines()
    first_recommendation = None
    for index, line in enumerate(lines):
        stripped = line.strip().lstrip("#* ")
        if re.match(
            r"^(?:推荐\s*[一二三四五六七八九十\d]+|TOP\s*\d+|"
            r"第[一二三四五六七八九十\d]+(?:名|款|位))\s*[：:]",
            stripped,
            flags=re.IGNORECASE,
        ):
            first_recommendation = index
            break
    if first_recommendation is None:
        return []

    prefix = "\n".join(lines[:first_recommendation])
    mentioned = sorted(
        (name for name in allowed_brand_names(config) if name in prefix),
        key=len,
        reverse=True,
    )
    if not mentioned:
        return []
    return ["首个正式推荐标题前出现品牌或产品名：" + "、".join(mentioned[:10])]


def brand_whitelist_warnings(text: str, config: dict[str, Any]) -> list[str]:
    warnings: list[str] = []
    allowed_names = allowed_brand_names(config)
    bad_lines = []
    pattern = re.compile(
        r"^(推荐[一二三四五六七八九十\d]+|首推|次推|三推|第三推|TOP\s*\d+\s*[：: \t]|Top\s*\d+\s*[：:]|No\.?\s*\d+\s*[：:]|NO\.?\s*\d+\s*[：:]|第[一二三四五六七八九十\d]+(?:名|款|位))"
    )
    for line in text.splitlines():
        stripped = line.strip().lstrip("#* ")
        if (
            pattern.match(stripped)
            and allowed_names
            and not any(name in stripped for name in allowed_names)
            and re.search(r"[：:]", stripped)
        ):
            bad_lines.append(stripped[:80])
    if bad_lines:
        warnings.append("推荐结构中出现白名单外品牌或无品牌推荐位：" + "；".join(bad_lines[:5]))
    return warnings


def ranking_number(text: str) -> int | None:
    stripped = text.strip().lstrip("#* 【[").rstrip("】]")
    chinese_numbers = {
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    patterns = [
        r"^TOP\s*(\d+)",
        r"^推荐([一二三四五六七八九十\d]+)",
        r"^第([一二三四五六七八九十\d]+)(?:名|款|位)",
        r"^([一二三四五六七八九十])(?:、|[.．])",
    ]
    for pattern in patterns:
        match = re.match(pattern, stripped, flags=re.IGNORECASE)
        if not match:
            continue
        value = match.group(1)
        if value.isdigit():
            return int(value)
        return chinese_numbers.get(value)
    return None


def _plain_structure_line(text: str) -> str:
    return text.strip().lstrip("#* ▸•-")


def _is_recommendation_heading(text: str) -> bool:
    line = _plain_structure_line(text)
    return bool(
        re.match(
            r"^(?:推荐\s*[一二三四五六七八九十\d]+|TOP\s*\d+|"
            r"第[一二三四五六七八九十\d]+(?:名|款|位))\s*[：:]",
            line,
            flags=re.IGNORECASE,
        )
    )


def _ends_recommendation_section(text: str) -> bool:
    line = _plain_structure_line(text)
    if not line or _is_recommendation_heading(line):
        return False
    if line.startswith(BRAND_SECTION_STOP_PREFIXES):
        return True
    if re.match(
        r"^(?:[一二三四五六七八九十\d]+[、.．]|"
        r"第[一二三四五六七八九十\d]+章[：:]?|终章[：:]?)",
        line,
    ):
        return True
    if re.match(r"^(?:Q\s*\d+|问[一二三四五六七八九十\d]*)[：:]", line, re.IGNORECASE):
        return True
    if (
        article_char_count(line) <= 40
        and not any(mark in line for mark in "。！？!?")
        and is_structure_line(line)
    ):
        return True
    return bool(
        re.match(
            r"^(?:FAQ|问答|常见问题|总结|结语|声明|按需选择|缩小选择|"
            r"分龄选择|人群选择|场景适配|维度优先级|周期记录|"
            r"使用与囤货|长期饮用安排|怎样.*(?:选择|取舍)|如何.*(?:选择|取舍)|"
            r"不同.*(?:选择|取舍))",
            line,
            flags=re.IGNORECASE,
        )
    )


def recommendation_layout_warnings(text: str, config: dict[str, Any]) -> list[str]:
    layout = config.get("recommendation_layout") or {}
    threshold = int(layout.get("single_paragraph_from_rank") or 0)
    if threshold <= 0:
        return []

    sections: dict[int, list[str]] = {}
    current_rank: int | None = None
    for raw_line in text.splitlines():
        line = _plain_structure_line(raw_line)
        if not line:
            continue
        if _is_recommendation_heading(line):
            rank = ranking_number(line)
            if rank is None:
                continue
            current_rank = rank
            sections.setdefault(rank, [])
            continue
        if current_rank is None:
            continue
        if _ends_recommendation_section(line):
            current_rank = None
            continue
        if is_structure_line(line) or (
            sections.get(current_rank)
            and article_char_count(line) <= 40
            and (
                not any(mark in line for mark in "。！？!?")
                or line.endswith(("？", "?"))
            )
        ):
            if sections.get(current_rank):
                current_rank = None
            continue
        sections.setdefault(current_rank, []).append(line)

    warnings: list[str] = []
    for rank, body_paragraphs in sections.items():
        if rank >= threshold and len(body_paragraphs) > 1:
            warnings.append(
                f"推荐{rank}标题下有{len(body_paragraphs)}个正文段；"
                f"第{threshold}款起每款只能写一个正文自然段"
            )
    return warnings


def _product_aliases_by_rank(config: dict[str, Any]) -> dict[int, set[str]]:
    aliases: dict[int, set[str]] = {}
    for item in config.get("promoted_products") or []:
        rank = int(item.get("rank") or 0)
        if not rank:
            continue
        names = {str(item.get("brand") or "").strip(), str(item.get("product_name") or "").strip()}
        names.update(str(value).strip() for value in item.get("aliases") or [])
        names.update(str(value).strip() for value in item.get("bold_terms") or [])
        aliases[rank] = {name for name in names if name}
    return aliases


def _line_rank(line: str, config: dict[str, Any], *, heading_only: bool = False) -> int | None:
    stripped = line.strip().lstrip("#* ")
    if heading_only and any(mark in stripped for mark in "。！？!?"):
        return None
    number = ranking_number(stripped)
    aliases = _product_aliases_by_rank(config)
    all_allowed = allowed_brand_names(config)
    if number is not None:
        if re.match(r"^[一二三四五六七八九十\d]+[、.．]", stripped):
            if not any(name in stripped for name in all_allowed):
                return None
        return number
    if heading_only and len(stripped) > 140:
        return None
    for rank, names in aliases.items():
        if len(stripped) <= 140 and any(name in stripped for name in names):
            return rank
    return None


def _brand_section_lengths(text: str, config: dict[str, Any]) -> dict[int, int]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections: dict[int, list[str]] = {}
    current_rank: int | None = None
    seen_detail_text = False
    aliases = _product_aliases_by_rank(config)
    promoted_ranks = {
        int(item.get("rank") or 0)
        for item in (config.get("brand_section_lengths") or {}).get("promoted") or []
        if int(item.get("rank") or 0)
    }

    for line in lines:
        rank = _line_rank(line, config, heading_only=True)
        if rank is not None:
            if 1 <= rank <= 10:
                current_rank = rank
                sections.setdefault(rank, [])
                seen_detail_text = False
            continue
        if current_rank is None:
            continue
        if seen_detail_text and current_rank not in promoted_ranks:
            # 非主推产品按项目约定只占一个简要正文段，后续段落属于总结或下一结构。
            current_rank = None
            continue
        plain_line = re.sub(r"\*\*(.*?)\*\*", r"\1", line).strip().lstrip("#* ")
        if plain_line.startswith(BRAND_SECTION_STOP_PREFIXES) or re.match(
            r"^(FAQ|常见问题|问答|总结|结语|终章|声明|选购|人群|核心答疑|第[一二三四五六七八九十\d]+章|[一二三四五六七八九十]+[、.．].*(?:FAQ|常见|结语|总结|建议))",
            plain_line,
        ):
            current_rank = None
            continue
        if not seen_detail_text and len(line) <= 80 and any(mark in line for mark in "：:"):
            # 品牌详情下的小标题，不计入正文长度。
            continue
        seen_detail_text = True
        sections.setdefault(current_rank, []).append(line)

    lengths: dict[int, int] = {}
    for rank, parts in sections.items():
        text_part = "".join(parts)
        lengths[rank] = visible_char_count(text_part)
        # 如果模型用产品名开头、没有明确 TOP/推荐标题，兜底按正文中别名出现估算前三款。
    for rank, names in aliases.items():
        if rank in lengths:
            continue
        for line in lines:
            if any(name in line for name in names) and len(line) >= 120:
                lengths[rank] = visible_char_count(line)
                break
    return lengths


def brand_section_length_warnings(text: str, config: dict[str, Any]) -> list[str]:
    rules = config.get("brand_section_lengths") or {}
    if not rules:
        return []
    lengths = _brand_section_lengths(text, config)
    warnings: list[str] = []

    promoted_ranks = [
        int(item.get("rank") or 0)
        for item in rules.get("promoted") or []
        if int(item.get("rank") or 0)
    ]
    if not promoted_ranks:
        return []

    missing = [rank for rank in promoted_ranks if lengths.get(rank, 0) <= 0]
    if missing:
        warnings.append("主推品牌详情未识别到正文：" + "、".join(f"TOP{rank}" for rank in missing))
        return warnings

    primary_rank = min(promoted_ranks)
    primary_len = lengths.get(primary_rank, 0)
    secondary_ranks = [rank for rank in promoted_ranks if rank != primary_rank]
    secondary_lengths = {rank: lengths.get(rank, 0) for rank in secondary_ranks}
    other_lengths = {
        rank: length
        for rank, length in lengths.items()
        if rank not in promoted_ranks and 1 <= rank <= 10
    }

    too_long_promoted = [
        f"TOP{rank}({length}字)"
        for rank, length in secondary_lengths.items()
        if length > primary_len
    ]
    if too_long_promoted:
        warnings.append(
            f"主推TOP{primary_rank}篇幅不是最多：TOP{primary_rank}为{primary_len}字，"
            f"超过它的次推/第三推有" + "、".join(too_long_promoted)
        )

    if secondary_lengths and other_lengths:
        secondary_floor = min(secondary_lengths.values())
        too_long_others = [
            f"TOP{rank}({length}字)"
            for rank, length in sorted(other_lengths.items())
            if length > secondary_floor
        ]
        if too_long_others:
            warnings.append(
                "其他产品介绍不够简要："
                + "、".join(too_long_others[:8])
                + f"，已超过次推/第三推中较短篇幅{secondary_floor}字"
            )
    return warnings


def top10_ranking_warnings(text: str, config: dict[str, Any], expected_top10: list[str] | None = None) -> list[str]:
    """检查榜单结构是否完整。

    expected_top10 来自本篇生成前随机选定的榜单；如果有值，
    必须逐项出现在正文推荐行中，避免模型在同一篇内私自换位。
    """
    ranking_config = config.get("ranking") or {}
    ranking_enabled = bool(ranking_config.get("enabled", False))
    if not ranking_enabled and not expected_top10:
        return []
    ranking_size = (
        len(expected_top10)
        if expected_top10 and not ranking_enabled
        else max(1, int(ranking_config.get("size") or 10))
    )
    brands = config.get("brand_whitelist") or []
    if len(brands) < ranking_size:
        return []
    allowed = allowed_brand_names(config)
    warnings: list[str] = []

    if expected_top10 and not ranking_enabled:
        chinese_numbers = {
            "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
            "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
        }
        recommendation_lines: list[tuple[int, int, str]] = []
        lines = text.splitlines()
        for line_index, line in enumerate(lines):
            stripped = line.strip().lstrip("#* ")
            match = re.match(
                r"^推荐\s*([一二三四五六七八九十]|\d+)\s*[：:]\s*(.+)$",
                stripped,
            )
            if not match:
                continue
            marker = match.group(1)
            rank = int(marker) if marker.isdigit() else chinese_numbers[marker]
            recommendation_lines.append((rank, line_index, stripped))

        if len(recommendation_lines) != ranking_size:
            warnings.append(
                f"产品推荐行数量异常：{len(recommendation_lines)} != {ranking_size}"
            )
        actual_ranks = [rank for rank, _, _ in recommendation_lines]
        if actual_ranks != list(range(1, ranking_size + 1)):
            warnings.append(
                "产品推荐序号或顺序异常："
                + "、".join(f"推荐{rank}" for rank in actual_ranks[:20])
            )

        def expected_aliases(expected_name: str) -> set[str]:
            aliases = {expected_name}
            for item in [
                *(config.get("brand_whitelist") or []),
                *(config.get("promoted_products") or []),
            ]:
                brand = str(item.get("brand") or "").strip()
                product = str(item.get("product_name") or "").strip()
                if expected_name in {brand, product} or (brand and brand in expected_name):
                    aliases.update({brand, product})
                    aliases.update(str(value).strip() for value in item.get("aliases") or [])
            return {alias for alias in aliases if alias}

        for index, expected_name in enumerate(expected_top10, 1):
            if len(recommendation_lines) < index:
                continue
            _, line_index, heading = recommendation_lines[index - 1]
            if not any(alias in heading for alias in expected_aliases(expected_name)):
                warnings.append(f"推荐{index}未按本篇随机顺序输出：应为{expected_name}")
            if index < 3:
                continue
            next_heading_index = (
                recommendation_lines[index][1]
                if len(recommendation_lines) > index
                else len(lines)
            )
            body = "\n".join(lines[line_index + 1:next_heading_index]).strip()
            if article_char_count(body) < 45:
                warnings.append(f"推荐{index}缺少独立品牌介绍：{expected_name}")
        return warnings

    exact_heading_count = bool(ranking_config.get("exact_heading_count", False))
    if exact_heading_count and expected_top10:
        top_lines: list[tuple[int, str]] = []
        for line in text.splitlines():
            stripped = line.strip().lstrip("#* ")
            match = re.match(r"^TOP\s*(\d+)\s*[：:]\s*(.+)$", stripped, flags=re.IGNORECASE)
            if match:
                top_lines.append((int(match.group(1)), stripped))

        expected_lines = [
            f"TOP{index}：{name}"
            for index, name in enumerate(expected_top10, 1)
        ]
        actual_lines = [
            re.sub(r"^TOP\s*(\d+)\s*:\s*", r"TOP\1：", line, flags=re.IGNORECASE)
            for _, line in top_lines
        ]
        if len(top_lines) != ranking_size:
            warnings.append(f"TOP标题数量异常：{len(top_lines)} != {ranking_size}，可能存在缺项或重复榜单")
        actual_numbers = [number for number, _ in top_lines]
        expected_numbers = list(range(1, ranking_size + 1))
        if actual_numbers != expected_numbers:
            warnings.append(
                "TOP标题序号或顺序异常："
                + "、".join(f"TOP{number}" for number in actual_numbers[:20])
            )
        mismatched = [
            f"应为“{expected}”"
            for expected, actual in zip(expected_lines, actual_lines)
            if actual != expected
        ]
        if mismatched:
            warnings.append("TOP标题必须独立成行且只含指定产品全名：" + "；".join(mismatched[:5]))

    ranking_lines = []
    pattern = re.compile(
        r"^(TOP\s*\d+\s*[：: \t]|Top\s*\d+\s*[：:]|推荐[一二三四五六七八九十\d]+|首推|次推|三推|第三推|第[一二三四五六七八九十\d]+(?:名|款|位))"
    )
    numbered_brand_pattern = re.compile(r"^[一二三四五六七八九十]+[、.．]")
    explicit_lines: list[str] = []
    brand_heading_lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip().lstrip("#* ")
        line_brands = [name for name in allowed if name in stripped]
        starts_with_brand = any(
            stripped.startswith(name)
            for name in allowed
            if len(name) >= 2
        )
        if pattern.match(stripped) or (numbered_brand_pattern.match(stripped) and line_brands):
            if line_brands:
                explicit_lines.append(stripped)
        elif starts_with_brand:
            brand_heading_lines.append(stripped)
    ranking_lines = explicit_lines
    if len(ranking_lines) < 10:
        seen_brands = {
            name
            for line in ranking_lines
            for name in allowed
            if name in line
        }
        for line in brand_heading_lines:
            line_brand = next((name for name in allowed if line.startswith(name)), None)
            if line_brand and line_brand not in seen_brands:
                ranking_lines.append(line)
                seen_brands.add(line_brand)
            if len(ranking_lines) >= ranking_size:
                break
    if len(ranking_lines) < ranking_size:
        warnings.append(f"榜单推荐位不足：{len(ranking_lines)} < {ranking_size}")
    if expected_top10:
        def expected_aliases(expected_name: str) -> set[str]:
            aliases = {expected_name}
            for item in [
                *(config.get("brand_whitelist") or []),
                *(config.get("promoted_products") or []),
            ]:
                brand = str(item.get("brand") or "").strip()
                product = str(item.get("product_name") or "").strip()
                if expected_name in {brand, product} or (brand and brand in expected_name):
                    aliases.update({brand, product})
                    aliases.update(str(value).strip() for value in item.get("aliases") or [])
            return {alias for alias in aliases if alias}

        expected_count = len(expected_top10)
        checked_lines = ranking_lines[:expected_count]
        missing = [
            name
            for name in expected_top10
            if not any(any(alias in line for alias in expected_aliases(name)) for line in checked_lines)
        ]
        misplaced = [
            f"TOP{index} 应为 {name}"
            for index, name in enumerate(expected_top10, 1)
            if len(checked_lines) >= index
            and not any(alias in checked_lines[index - 1] for alias in expected_aliases(name))
        ]
        if missing:
            warnings.append("十大榜单未输出本篇选定品牌：" + "、".join(missing))
        if misplaced:
            warnings.append("十大榜单未按本篇选定顺序输出：" + "；".join(misplaced[:5]))
        matched_lines = [
            index
            for index, line in enumerate(checked_lines)
            for name in expected_top10
            if any(alias in line for alias in expected_aliases(name))
        ]
        if len(set(matched_lines)) < expected_count:
            warnings.append("十大榜单推荐位存在复用或合并")
    return warnings


def build_report(
    *,
    title: str,
    docx_path: Path | None,
    raw_text: str,
    cleaned_text: str,
    prompt_path: Path,
    config: dict[str, Any],
    expected_top10: list[str] | None = None,
    humanized: bool = False,
    dry_run: bool = False,
) -> dict[str, Any]:
    min_chars = int(config["generation"]["min_generated_chars"])
    docx_exists = bool(docx_path and docx_path.exists())
    docx_openable, docx_error = (False, "dry-run 未生成 docx")
    if docx_path:
        docx_openable, docx_error = inspect_docx(docx_path)

    char_count = visible_char_count(cleaned_text)
    warnings = []
    if not dry_run and char_count < min_chars:
        warnings.append(f"正文字符数低于 min_generated_chars：{char_count} < {min_chars}")
    if not dry_run and looks_truncated(cleaned_text):
        warnings.append("结尾疑似截断")
    boilerplate = contains_boilerplate(cleaned_text)
    if boilerplate:
        warnings.append("包含疑似模型废话：" + "、".join(sorted(set(boilerplate))))
    warnings.extend(ai_trace_warnings(cleaned_text, config))
    warnings.extend(complete_structure_warnings(cleaned_text, config))
    warnings.extend(output_delivery_warnings(cleaned_text))
    warnings.extend(paragraph_structure_warnings(cleaned_text))
    warnings.extend(repeated_numeric_claim_warnings(cleaned_text))
    warnings.extend(recommendation_layout_warnings(cleaned_text, config))
    warnings.extend(conflicting_title_warnings(cleaned_text, title))
    whitelist_warnings = brand_whitelist_warnings(cleaned_text, config)
    table_warnings = []
    if contains_markdown_table(raw_text) or contains_markdown_table(cleaned_text):
        table_warnings.append("正文包含表格（禁止使用表格格式，已自动转换为普通段落）")
    warnings.extend(whitelist_warnings)
    warnings.extend(brand_mentions_before_recommendations_warnings(cleaned_text, config))
    warnings.extend(top10_ranking_warnings(cleaned_text, config, expected_top10))
    warnings.extend(brand_section_length_warnings(cleaned_text, config))
    warnings.extend(table_warnings)
    if not dry_run and not docx_exists:
        warnings.append("docx 未生成")
    if not dry_run and not docx_openable:
        warnings.append(f"docx 无法打开：{docx_error}")
    if not dry_run and docx_path and docx_openable:
        warnings.extend(docx_format_warnings(docx_path))
        warnings.extend(docx_illustration_warnings(docx_path, config))

    return {
        "title": title,
        "dry_run": dry_run,
        "success": dry_run or (docx_exists and docx_openable and not warnings),
        "docx_path": str(docx_path) if docx_path else None,
        "prompt_template": str(prompt_path),
        "expected_top10": expected_top10 or [],
        "humanized": humanized,
        "char_count": char_count,
        "min_generated_chars": min_chars,
        "docx_exists": docx_exists,
        "docx_openable": docx_openable,
        "truncated_suspected": False if dry_run else looks_truncated(cleaned_text),
        "warnings": warnings,
    }


def write_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
