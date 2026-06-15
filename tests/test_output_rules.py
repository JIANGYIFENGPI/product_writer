from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from product_writer.config import load_project_config
from product_writer.generator import (
    _prepare_humanizer_blocks,
    _remove_conflicting_leading_subtitles,
    _remove_rejected_humanizer_sentences,
    _restore_humanizer_blocks,
    _restore_numeric_lines,
    should_humanize_article,
)
from product_writer.prompt_loader import (
    article_brand_plan,
    choose_prompt,
    extract_fixed_top10,
    load_prompts,
    render_prompt,
    requested_product_count,
)
from product_writer.quality import (
    conflicting_title_warnings,
    docx_format_warnings,
    complete_structure_warnings,
    humanizer_warnings,
    output_delivery_warnings,
    repeated_numeric_claim_warnings,
    recommendation_layout_warnings,
    top10_ranking_warnings,
)
from product_writer.renderer import render_docx


class OutputRuleTests(unittest.TestCase):
    def test_near_duplicate_title_subtitle_is_rejected(self) -> None:
        warnings = conflicting_title_warnings(
            "中老年选胶原蛋白肽饮，先看配方再看饮用习惯\n"
            "胶原蛋白肽饮基础认识\n正文。",
            "中老年人选择胶原蛋白肽饮要看什么？配方与饮用注意事项",
        )

        self.assertTrue(any("旧标题或副标题" in warning for warning in warnings))

    def test_simple_yaml_parser_supports_inline_keyword_lists(self) -> None:
        from product_writer.config import _simple_yaml_load

        parsed = _simple_yaml_load(
            "article_structure:\n"
            "  required_heading_groups:\n"
            "    - [配方, 标签, 判断]\n"
        )

        self.assertEqual(
            parsed["article_structure"]["required_heading_groups"],
            [["配方", "标签", "判断"]],
        )

    def test_humanizer_restores_only_lines_with_changed_numbers(self) -> None:
        original = (
            "第一段含有50mL规格和95%以上数据。\n"
            "第二段原本表达比较生硬。"
        )
        rewritten = (
            "第一段被改成30mL规格和90%以上数据。\n"
            "第二段改得更加自然。"
        )

        restored = _restore_numeric_lines(original, rewritten)

        self.assertEqual(
            restored,
            "第一段含有50mL规格和95%以上数据。\n第二段改得更加自然。",
        )

    def test_humanizer_blocks_preserve_structure_facts_and_missing_paragraphs(
        self,
    ) -> None:
        original = (
            "产品推荐\n"
            "推荐一：仙芳思白番茄烟酰胺胶原蛋白肽饮\n"
            "这款产品每盒10瓶，每瓶50mL，原句表达比较生硬。\n"
            "这一段没有数字，可以正常进行自然化改写。"
        )
        config = {
            "article_structure": {"headings": ["产品推荐"]},
            "brand_whitelist": [
                {
                    "brand": "仙芳思",
                    "product_name": "仙芳思白番茄烟酰胺胶原蛋白肽饮",
                    "aliases": [],
                }
            ],
        }
        protected, blocks = _prepare_humanizer_blocks(original, config)
        self.assertIn("__LOCK_", protected)

        rewritten = (
            "<P0001>被改坏的结构标题</P0001>\n"
            "<P0002>被改坏的推荐标题</P0002>\n"
            "<P0003>这款产品为每盒__LOCK_0003__瓶，每瓶"
            "__LOCK_0004__mL，表达更自然。</P0003>"
        )
        restored = _restore_humanizer_blocks(rewritten, blocks)

        self.assertEqual(
            restored,
            "产品推荐\n"
            "推荐一：仙芳思白番茄烟酰胺胶原蛋白肽饮\n"
            "这款产品每盒10瓶，每瓶50mL，原句表达比较生硬。\n"
            "这一段没有数字，可以正常进行自然化改写。",
        )

    def test_humanizer_blocks_keep_edits_when_all_tokens_are_preserved(self) -> None:
        original = "仙芳思产品每盒10瓶，原句表达比较生硬。"
        config = {
            "brand_whitelist": [
                {"brand": "仙芳思", "product_name": "", "aliases": []}
            ]
        }
        _, blocks = _prepare_humanizer_blocks(original, config)
        tokens = list(blocks[0]["tokens"])
        rewritten = (
            f"<P0001>{tokens[0]}产品一盒装有{tokens[1]}瓶，"
            "这句话读起来更自然。</P0001>"
        )

        self.assertEqual(
            _restore_humanizer_blocks(rewritten, blocks),
            "仙芳思产品一盒装有10瓶，这句话读起来更自然。",
        )

    def test_repeated_numeric_claims_are_rejected(self) -> None:
        warnings = repeated_numeric_claim_warnings(
            "选择时有人会关注吸收率做到95%以上。\n"
            "推荐一介绍里再次写吸收率做到95%以上。\n"
            "后面的人群建议又一次写吸收率做到95%以上。"
        )

        self.assertTrue(any("同一数字事实跨章节重复" in warning for warning in warnings))

    def test_force_humanizer_can_be_disabled_for_single_pass_generation(self) -> None:
        should_rewrite, findings = should_humanize_article(
            "这是一段自然、简洁且没有固定模板词的正文。",
            {
                "humanizer": {
                    "enabled": True,
                    "force_rewrite": False,
                }
            },
        )

        self.assertFalse(should_rewrite)
        self.assertEqual(findings, [])

    def test_prompt_contains_program_delivery_contract(self) -> None:
        config = {
            "project": {"name": "测试产品"},
            "generation": {
                "min_generated_chars": 3500,
                "target_generated_chars": 4000,
            },
            "features": {"images": False},
            "brand_whitelist": [
                {"brand": "测试品牌", "product_name": "测试产品"},
            ],
            "promoted_products": [],
            "ranking": {"enabled": False},
        }

        prompt = render_prompt("围绕 {title} 写作。", "测试标题", config)

        self.assertIn("不要自行添加星号加粗", prompt)
        self.assertIn("程序会统一生成宋体 12pt", prompt)
        self.assertIn("当前阶段不插图", prompt)
        self.assertIn("不得输出“此处插图”", prompt)
        self.assertNotIn("虚拟品牌", prompt)
        self.assertIn("测试品牌", prompt)

    def test_image_enabled_prompt_delegates_images_to_program(self) -> None:
        config = {
            "project": {"name": "测试产品"},
            "generation": {
                "min_generated_chars": 3500,
                "target_generated_chars": 4000,
            },
            "features": {"images": True},
            "brand_whitelist": [
                {"brand": "测试品牌", "product_name": "测试产品"},
            ],
            "promoted_products": [],
            "ranking": {"enabled": False},
        }

        prompt = render_prompt("围绕 {title} 写作。", "测试标题", config)

        self.assertIn("图片由程序从项目素材槽位自动选择并插入", prompt)
        self.assertNotIn("当前阶段不插图", prompt)

    def test_renderer_output_passes_format_checks(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "sample.docx"
            render_docx(
                "测试标题",
                "一、选购建议\n正文内容包含测试关键词。",
                output_path,
                config,
                ["测试关键词"],
            )

            self.assertEqual(docx_format_warnings(output_path), [])
            document = Document(output_path)
            self.assertTrue(all(run.bold for run in document.paragraphs[0].runs))
            self.assertTrue(all(run.bold for run in document.paragraphs[1].runs))
            keyword_runs = [
                run
                for run in document.paragraphs[2].runs
                if run.text == "测试关键词"
            ]
            self.assertEqual(len(keyword_runs), 1)
            self.assertTrue(keyword_runs[0].bold)

    def test_renderer_bolds_configured_headings_but_not_full_faq_questions(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "article_structure": {
                "headings": [
                    "配方与标签怎么看",
                    "常见问题答疑",
                    "选购总结",
                ]
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "headings.docx"
            render_docx(
                "测试标题",
                "配方与标签怎么看\n正文内容。\n"
                "常见问题答疑\n一天什么时候喝比较合适？\n回答内容。\n"
                "选购总结\n总结内容。",
                output_path,
                config,
                [],
            )

            document = Document(output_path)
            paragraph_by_text = {
                paragraph.text: paragraph
                for paragraph in document.paragraphs
                if paragraph.text
            }
            for heading in (
                "配方与标签怎么看",
                "常见问题答疑",
                "选购总结",
            ):
                self.assertTrue(
                    all(run.bold for run in paragraph_by_text[heading].runs)
                )
            self.assertFalse(
                any(
                    bool(run.bold)
                    for run in paragraph_by_text["一天什么时候喝比较合适？"].runs
                )
            )

    def test_renderer_bolds_product_dimension_labels(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "labels.docx"
            render_docx(
                "测试标题",
                "原料与专利：这里介绍原料来源和专利信息。\n"
                "销量复购：这里介绍销量和复购率数据。",
                output_path,
                config,
                [],
            )

            document = Document(output_path)
            for paragraph, expected_label in zip(
                document.paragraphs[1:],
                ("原料与专利：", "销量复购："),
            ):
                self.assertEqual(paragraph.runs[0].text, expected_label)
                self.assertTrue(paragraph.runs[0].bold)
                self.assertFalse(paragraph.runs[1].bold)

    def test_renderer_bolds_any_short_leading_colon_label(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "generic-label.docx"
            render_docx(
                "测试标题",
                "自定义观察项：这里是不在预设词库中的正文内容。",
                output_path,
                config,
                [],
            )

            paragraph = Document(output_path).paragraphs[1]
            self.assertEqual(paragraph.runs[0].text, "自定义观察项：")
            self.assertTrue(paragraph.runs[0].bold)
            self.assertFalse(paragraph.runs[1].bold)

    def test_renderer_does_not_bold_full_numbered_explanation(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "numbered-explanation.docx"
            render_docx(
                "测试标题",
                "1. 分子量信息：查看检测口径、区间和批次信息，"
                "不要把单一数字直接等同于产品优劣。",
                output_path,
                config,
                [],
            )

            paragraph = Document(output_path).paragraphs[1]
            self.assertEqual(paragraph.runs[0].text, "1. 分子量信息：")
            self.assertTrue(paragraph.runs[0].bold)
            self.assertTrue(any(not run.bold for run in paragraph.runs[1:]))

    def test_renderer_bolds_product_list_heading(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "product-list-heading.docx"
            render_docx(
                "测试标题",
                "十款产品逐一看\n推荐一：测试产品\n这里是产品正文。",
                output_path,
                config,
                [],
            )

            paragraph = Document(output_path).paragraphs[1]
            self.assertEqual(paragraph.text, "十款产品逐一看")
            self.assertTrue(all(run.bold for run in paragraph.runs))

    def test_renderer_bolds_full_faq_section_heading(self) -> None:
        config = {
            "features": {
                "bold_structure": True,
                "bold_terms": True,
            },
            "promoted_products": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "faq-heading.docx"
            render_docx(
                "测试标题",
                "白番茄烟酰胺高频FAQ\n什么时候饮用比较合适？\n回答正文。",
                output_path,
                config,
                ["白番茄烟酰胺"],
            )

            paragraph = Document(output_path).paragraphs[1]
            self.assertEqual(paragraph.text, "白番茄烟酰胺高频FAQ")
            self.assertTrue(all(run.bold for run in paragraph.runs))

    def test_format_checks_reject_heading_and_wrong_font(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "bad.docx"
            document = Document()
            paragraph = document.add_heading("错误标题", level=1)
            paragraph.runs[0].font.name = "微软雅黑"
            document.save(output_path)

            warnings = docx_format_warnings(output_path)

            self.assertTrue(any("非普通段落样式" in warning for warning in warnings))
            self.assertTrue(any("宋体 12pt" in warning for warning in warnings))

    def test_output_delivery_checks_image_placeholders(self) -> None:
        warnings = output_delivery_warnings("正文内容。\n此处插图：产品展示")
        self.assertTrue(any("图片占位" in warning for warning in warnings))

    def test_blank_line_is_not_treated_as_indentation(self) -> None:
        self.assertEqual(output_delivery_warnings("第一段。\n\n第二段。"), [])
        warnings = output_delivery_warnings("第一段。\n　第二段。")
        self.assertTrue(any("行首空格" in warning for warning in warnings))

    def test_humanizer_detects_repeated_template_openings(self) -> None:
        text = "\n".join(
            [
                "购买前需要先查看配料表，再结合自己的饮食安排作判断。" * 2,
                "购买前需要核对实物标签，不能只根据包装正面作判断。" * 2,
                "购买前需要留意建议食用量，避免把规格等同于单次用量。" * 2,
                "购买前需要确认储存方式，并按开封后的说明及时饮用。" * 2,
                "购买前需要比较单瓶容量和整盒数量，再估算一段时间的使用成本。" * 2,
                "购买前需要结合每日建议量计算每盒可用天数，避免只比较整盒价格。" * 2,
                "购买前需要留意口味和便携方式，确认是否符合自己的饮用习惯。" * 2,
                "消费者需要核对配料表和生产许可证，信息不全时应谨慎。" * 2,
            ]
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("模板化AI腔" in warning for warning in warnings))
        self.assertTrue(any("相同开头" in warning for warning in warnings))

    def test_humanizer_detects_repeated_product_intro_template(self) -> None:
        text = "\n".join(
            [
                "选择这款产品时，可以先查看包装名称和建议用量，再结合自己的预算判断。" * 2,
                "选择这款产品时，可以先核对配料顺序和规格，再决定是否适合长期安排。" * 2,
                "选择这款产品时，可以先确认储存方式和开封要求，避免只看宣传名称。" * 2,
            ]
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("产品介绍使用相同套话" in warning for warning in warnings))

    def test_humanizer_detects_semantically_repeated_adjacent_paragraphs(self) -> None:
        text = (
            "把一瓶胶原蛋白肽饮翻到背面，每个人关注的位置不同。有人找胶原蛋白肽，"
            "有人看烟酰胺，也有人直接计算营养成分表里的蛋白质含量。\n"
            "拿到一瓶胶原蛋白肽饮后翻看背面，不同人关注的位置并不一样。有人先找"
            "胶原蛋白肽，有人看有没有烟酰胺，还有人计算蛋白质含量。"
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("近义重复段落" in warning for warning in warnings))

    def test_humanizer_detects_dense_abstract_ai_transitions(self) -> None:
        text = (
            "真正关键的不是包装，而是配料。本质上，这是选择路径的问题。"
            "这意味着数字可以成为判断信息完整度的试金石。"
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("抽象判断和转折套话过多" in warning for warning in warnings))

    def test_humanizer_allows_normal_product_positioning_phrase(self) -> None:
        warnings = humanizer_warnings(
            "这款产品定位偏向日常即饮，包装和规格适合放在办公室。"
        )

        self.assertFalse(any("标签核对式AI套话" in warning for warning in warnings))

    def test_humanizer_only_rejects_heavily_repeated_openings(self) -> None:
        four_paragraphs = "\n".join(
            f"胶原蛋白肽饮第{index}段提供不同的配方、规格和饮用场景信息，"
            "每段内容都足够长，可以用于检查开头重复阈值，同时补充包装与日常安排。"
            for index in range(1, 5)
        )
        seven_paragraphs = four_paragraphs + "\n" + "\n".join(
            f"胶原蛋白肽饮第{index}段继续提供不同的包装、风味和日常安排信息，"
            "用于确认严重重复仍会被自动质检拦截，并保留足够的正文长度。"
            for index in range(5, 8)
        )

        self.assertFalse(
            any("多段使用相同开头" in warning for warning in humanizer_warnings(four_paragraphs))
        )
        self.assertTrue(
            any("多段使用相同开头" in warning for warning in humanizer_warnings(seven_paragraphs))
        )

    def test_humanizer_removes_rejected_sentences_and_ai_prefixes(self) -> None:
        cleaned = _remove_rejected_humanizer_sentences(
            "序号不代表产品优劣，仅作选购信息参考。\n"
            "理清这十二个维度之后，我们再进入今年的年度产品清单。"
            "本篇一共收录十款，排序不代表功效高低或商业排名，"
            "只是年度信息整理的一种呈现方式。\n"
            "数据显示，这款产品规格为50mL。\n"
            "本质上，这意味着真正关键的是按实际需求选择。\n"
            "这样写的目的是让读者核对标签和信息完整度。"
        )

        self.assertNotIn("不代表产品优劣", cleaned)
        self.assertNotIn("本篇一共收录", cleaned)
        self.assertNotIn("排序不代表功效高低", cleaned)
        self.assertNotIn("年度信息整理的一种呈现方式", cleaned)
        self.assertNotIn("这篇文章", cleaned)
        self.assertNotIn("这篇内容", cleaned)
        self.assertNotIn("数据显示", cleaned)
        self.assertNotIn("本质上", cleaned)
        self.assertNotIn("这意味着", cleaned)
        self.assertIn("这款产品规格为50mL", cleaned)
        self.assertIn("便于读者查看配料和规格和可查信息", cleaned)

    def test_humanizer_removes_conflicting_leading_subtitle(self) -> None:
        title = "胶原蛋白肽饮怎么比较检测与认证信息？2026选购参考"
        text = (
            "胶原蛋白肽饮的检测和认证信息，怎么比着看？2026选购参考\n"
            "胶原蛋白肽饮基础认识\n"
            "这里是正常正文。"
        )

        self.assertEqual(
            _remove_conflicting_leading_subtitles(text, title),
            "胶原蛋白肽饮基础认识\n这里是正常正文。",
        )

    def test_humanizer_rejects_unsupported_health_inference(self) -> None:
        warnings = humanizer_warnings(
            "液态产品胃排空速度比片剂快，还能降低肠胃的工作负荷。"
        )

        self.assertTrue(any("医学或效果推断" in warning for warning in warnings))

    def test_humanizer_rejects_sales_data_causal_claim(self) -> None:
        warnings = humanizer_warnings(
            "这款产品复购率更能体现饮用体验，销量反映市场接受度。"
        )

        self.assertTrue(any("销量、复购或认证数据" in warning for warning in warnings))

    def test_humanizer_rejects_label_disclaimer_phrasing(self) -> None:
        warnings = humanizer_warnings(
            "这款产品目前只有包装名称，购买时还需核对实物标签，以实物为准。"
        )

        self.assertTrue(any("标签核对式AI套话" in warning for warning in warnings))

    def test_humanizer_rejects_prompt_rule_echo_as_product_analysis(self) -> None:
        text = (
            "仙芳思的白番茄只在这款产品中出现，不构成整个胶原蛋白肽品类的必选项。"
            "它与其他配方形成两条不同路线，选择哪条路线取决于自己日常饮食中"
            "哪类营养相对不足。"
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("标签核对式AI套话" in warning for warning in warnings))

    def test_humanizer_rejects_editorial_language_inside_product_intro(self) -> None:
        text = (
            "推荐三：测试胶原蛋白肽饮\n"
            "它的资料没有堆叠很多复配成分，标签语言越简短，"
            "越要避免自行补充没有提供的结论。\n"
            "按需选择\n"
            "这里是正常的通用选择建议。"
        )

        warnings = humanizer_warnings(text)

        self.assertTrue(any("内部资料或编辑规则口吻" in warning for warning in warnings))

    def test_humanizer_rejects_ranking_disclaimer(self) -> None:
        warnings = humanizer_warnings(
            "推荐一：甲产品\n序号不代表市场排名或优劣，以下内容仅作选购信息参考。",
            {"recommendation_layout": {"reject_ranking_disclaimer": True}},
        )

        self.assertTrue(any("标签核对式AI套话" in warning for warning in warnings))

    def test_humanizer_allows_ranking_disclaimer_when_project_rule_is_disabled(
        self,
    ) -> None:
        warnings = humanizer_warnings(
            "推荐一：甲产品\n序号不代表市场排名或优劣，以下内容仅作选购信息参考。",
            {"recommendation_layout": {"reject_ranking_disclaimer": False}},
        )

        self.assertFalse(any("标签核对式AI套话" in warning for warning in warnings))

    def test_recommendation_layout_rejects_multiple_body_paragraphs_from_rank_three(
        self,
    ) -> None:
        text = (
            "推荐三：甲产品\n"
            "第一段介绍产品的配方、口感和适合的饮用场景，内容足够形成正文。\n"
            "第二段又补充包装和使用体验，因此已经拆成了两个正文自然段。\n"
            "推荐四：乙产品\n"
            "这一款只保留一个完整正文自然段，用自然衔接写清产品特点和适用场景。"
        )

        warnings = recommendation_layout_warnings(
            text,
            {"recommendation_layout": {"single_paragraph_from_rank": 3}},
        )

        self.assertTrue(any("推荐3标题下有2个正文段" in warning for warning in warnings))

    def test_recommendation_layout_stops_at_numbered_section_heading(self) -> None:
        text = (
            "推荐三：甲产品\n"
            "这一款只保留一个完整正文自然段，用自然衔接写清产品特点和适用场景。\n"
            "六、分龄选择与问答\n"
            "这里是分龄建议，不属于推荐三的产品介绍。\n"
            "问题一：应该怎么安排？\n"
            "这里是问题回答。"
        )

        warnings = recommendation_layout_warnings(
            text,
            {"recommendation_layout": {"single_paragraph_from_rank": 3}},
        )

        self.assertEqual(warnings, [])

    def test_recommendation_layout_stops_at_unlisted_structure_heading(self) -> None:
        text = (
            "推荐十：甲产品\n"
            "这一款只保留一个完整正文自然段，用自然衔接写清产品特点和适用场景。\n"
            "购买时还要考虑哪些现实条件\n"
            "这里讨论预算、储存和长期安排，不属于推荐十的产品介绍。\n"
            "最后的选择建议\n"
            "这里是全文收束。"
        )

        warnings = recommendation_layout_warnings(
            text,
            {"recommendation_layout": {"single_paragraph_from_rank": 3}},
        )

        self.assertEqual(warnings, [])

    def test_recommendation_layout_stops_at_short_heading_and_qa(self) -> None:
        text = (
            "推荐十：甲产品\n"
            "这一款只保留一个完整正文自然段，用自然衔接写清产品特点和适用场景。\n"
            "长期安排怎么做\n"
            "这里讨论预算和储存，不属于推荐十的产品介绍。\n"
            "Q1：空腹可以喝吗？\n"
            "这里是问题回答。"
        )

        warnings = recommendation_layout_warnings(
            text,
            {"recommendation_layout": {"single_paragraph_from_rank": 3}},
        )

        self.assertEqual(warnings, [])

    def test_product_editorial_scan_stops_at_numbered_section_heading(self) -> None:
        text = (
            "推荐三：甲产品\n"
            "这一款介绍配方、规格和口感，不包含编辑过程说明。\n"
            "六、常见问题与结语\n"
            "已有资料说明这一问题需要单独讨论。"
        )

        warnings = humanizer_warnings(text)

        self.assertFalse(any("内部资料或编辑规则口吻" in warning for warning in warnings))

    def test_prompt_selection_matches_title_angle(self) -> None:
        prompts = [
            (Path("05_测评维度.txt"), "测评"),
            (Path("08_胶原深度.txt"), "胶原"),
            (Path("01_通用指南.txt"), "通用"),
        ]
        config = {
            "project": {
                "prompt_selection_rules": [
                    {
                        "keywords": ["胶原蛋白肽怎么选"],
                        "prompt_prefixes": ["08_"],
                    },
                    {
                        "keywords": ["实测"],
                        "prompt_prefixes": ["05_"],
                    },
                ],
                "default_prompt_prefixes": ["01_"],
            }
        }

        self.assertEqual(
            choose_prompt(prompts, "胶原蛋白肽怎么选？", config)[0].name,
            "08_胶原深度.txt",
        )
        self.assertEqual(
            choose_prompt(
                [
                    (Path("08_胶原深度.txt"), "胶原"),
                    (Path("09_胶原周期.txt"), "周期"),
                ],
                "胶原蛋白肽需要喝多久？",
                {
                    "project": {
                        "prompt_selection_rules": [
                            {
                                "keywords": ["需要喝多久"],
                                "prompt_prefixes": ["09_"],
                            }
                        ]
                    }
                },
            )[0].name,
            "09_胶原周期.txt",
        )
        self.assertEqual(
            choose_prompt(prompts, "十款产品实测参考", config)[0].name,
            "05_测评维度.txt",
        )

    def test_complete_structure_requires_independent_sections(self) -> None:
        config = {
            "article_structure": {
                "enabled": True,
                "required_heading_groups": [
                    ["基础", "科普"],
                    ["选购", "判断"],
                    ["常见问题", "问答"],
                    ["结语", "总结"],
                ],
            }
        }
        warnings = complete_structure_warnings(
            "基础知识\n正文。\n选购判断\n正文。\n常见问题\n问题与回答。\n结语\n正文。",
            config,
        )
        self.assertEqual(warnings, [])
        warnings = complete_structure_warnings("基础知识\n正文。\n结语\n正文。", config)
        self.assertTrue(any("选购/判断" in warning for warning in warnings))

    def test_complete_structure_can_require_minimum_heading_count(self) -> None:
        warnings = complete_structure_warnings(
            "基础认识\n正文。\n产品推荐\n正文。\n选购总结\n正文。",
            {
                "article_structure": {
                    "enabled": True,
                    "minimum_heading_count": 5,
                    "required_heading_groups": [["推荐"], ["总结"]],
                }
            },
        )

        self.assertTrue(any("独立结构标题不足" in warning for warning in warnings))

    def test_project_requires_non_empty_brand_whitelist(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            project = root / "projects" / "sample"
            project.mkdir(parents=True)
            (project / "project.yaml").write_text(
                "project:\n  name: 测试产品\n",
                encoding="utf-8",
            )
            (project / "brands.txt").write_text("", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "品牌白名单为空"):
                load_project_config(root, "sample")

    def test_promoted_product_must_exist_in_whitelist(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            project = root / "projects" / "sample"
            project.mkdir(parents=True)
            (project / "project.yaml").write_text(
                "project:\n"
                "  name: 测试产品\n"
                "promoted_products:\n"
                "  - rank: 1\n"
                "    brand: 测试品牌\n"
                "    product_name: 未登记产品\n",
                encoding="utf-8",
            )
            (project / "brands.txt").write_text(
                "测试品牌 | 白名单产品\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "必须完整写入 brands.txt"):
                load_project_config(root, "sample")

    def test_product_count_defaults_to_ten_and_honors_title_count(self) -> None:
        self.assertEqual(requested_product_count("白番茄产品怎么选"), 10)
        self.assertEqual(requested_product_count("六款白番茄产品参考"), 6)
        self.assertEqual(requested_product_count("10款口服液清单"), 10)

    def test_five_product_title_overrides_ten_product_template_wording(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")
        prompt = render_prompt(
            "专用结构中原本要求十款产品和推荐三至十。",
            "白番茄烟酰胺品牌怎么选？2026五款产品推荐",
            config,
        )

        self.assertIn("【本篇推荐数量最高优先级】", prompt)
        self.assertIn("本篇只能输出推荐一至推荐五，共5款", prompt)
        self.assertIn("禁止输出推荐6及其后的产品", prompt)

    def test_brand_intro_plan_keeps_promoted_products_and_randomizes_others(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")
        candidates = [
            item
            for item in config["brand_whitelist"]
            if item["brand"] not in {"仙芳思", "赤大师"}
        ]
        first_selection = [dict(item) for item in candidates[:8]]
        second_selection = [dict(item) for item in reversed(candidates[:8])]

        with patch(
            "product_writer.prompt_loader.random.sample",
            side_effect=[first_selection, second_selection],
        ):
            first = article_brand_plan("十款胶原蛋白肽饮参考", config)
            second = article_brand_plan("十款胶原蛋白肽饮参考", config)

        self.assertEqual([item["brand"] for item in first[:2]], ["仙芳思", "赤大师"])
        self.assertEqual([item["brand"] for item in second[:2]], ["仙芳思", "赤大师"])
        self.assertNotEqual(
            [item["brand"] for item in first[2:]],
            [item["brand"] for item in second[2:]],
        )
        self.assertEqual([item["rank"] for item in first], list(range(1, 11)))
        self.assertTrue(all(item.get("profile") for item in first[2:]))
        self.assertTrue(all(not item.get("writing_mode") for item in first))

    def test_rendered_prompt_records_random_recommendation_order(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")
        config["project"]["require_other_brand_details"] = False
        prompt = render_prompt("围绕 {title} 写作。", "十款胶原蛋白肽饮参考", config)
        expected = extract_fixed_top10(prompt)

        self.assertEqual(len(expected), 10)
        self.assertTrue(expected[0].startswith("仙芳思"))
        self.assertTrue(expected[1].startswith("赤大师"))
        self.assertEqual(len(set(expected)), 10)
        self.assertIn("本篇产品推荐顺序", prompt)
        self.assertNotIn("驼奶粉基础认识", prompt)
        self.assertNotIn("【全文必备结构标题】", prompt)
        self.assertNotIn("本款组织方式：", prompt)

    def test_white_tomato_project_has_complete_updated_product_profiles(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")

        self.assertEqual(config["project"]["name"], "\u767d\u756a\u8304\u70df\u9170\u80fa")
        self.assertEqual(len(config["brand_whitelist"]), 32)
        self.assertEqual(len(config["brand_profiles"]), 30)
        self.assertEqual(
            [item["brand"] for item in config["promoted_products"]],
            ["仙芳思", "赤大师"],
        )
        self.assertEqual(
            {
                (item["brand"], item["product_name"])
                for item in config["brand_whitelist"]
                if item["brand"] not in {"仙芳思", "赤大师"}
            },
            {
                (item["brand"], item["product_name"])
                for item in config["brand_profiles"]
            },
        )
        self.assertTrue(config["features"]["images"])
        self.assertEqual(
            config["illustrations"]["neutral_placement_strategy"],
            "section_boundary_with_spacing",
        )
        self.assertEqual(
            config["illustrations"]["neutral_body_paragraph_min_chars"],
            60,
        )
        self.assertEqual(
            config["illustrations"]["minimum_body_paragraphs_between_images"],
            3,
        )
        self.assertNotIn(
            "cover_after_body_paragraphs",
            config["illustrations"],
        )
        self.assertEqual(config["illustrations"]["required_count"], 4)
        self.assertEqual(
            config["illustrations"]["extra_images_by_promoted_rank"],
            {"1": 1},
        )
        self.assertEqual(
            config["brand_section_lengths"]["promoted"][0],
            {"rank": 1, "min_chars": 800, "max_chars": 1100},
        )
        self.assertGreater(
            len(config["illustrations"]["neutral_image_keywords"]),
            10,
        )
        self.assertIn("\u767d\u756a\u8304", str(config["brand_profiles"]))
        self.assertIn("\u4ed9\u82b3\u601d", str(config["brand_whitelist"]))
        self.assertIn("\u8d64\u5927\u5e08", str(config["brand_whitelist"]))
        self.assertNotIn("\u5a07\u5c0f\u989c", str(config["brand_whitelist"]))
        self.assertTrue(config["article_structure"]["enabled"])
        self.assertFalse(config["article_structure"]["inject_prompt_headings"])
        self.assertTrue(
            any(
                "总结" in heading
                for heading in config["article_structure"]["headings"]
            )
        )

    def test_white_tomato_prompts_inject_eight_other_brand_profiles(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")
        project = root / "projects" / "jiaoyuandanbaitai"
        prompts = load_prompts(project, config)

        self.assertEqual(len(prompts), 10)
        for _, template in prompts:
            rendered = render_prompt(
                template,
                "2026\u80f6\u539f\u86cb\u767d\u80bd\u996e\u600e\u4e48\u9009\uff1f"
                "\u5341\u6b3e\u4ea7\u54c1\u53c2\u8003",
                config,
            )
            self.assertEqual(
                rendered.count(
                    "\u5185\u90e8\u7d20\u6750\uff08\u53ea\u4f9b\u6539\u5199\uff0c"
                    "\u7981\u6b62\u7167\u6284\u5b57\u6bb5\u540d\uff09"
                ),
                8,
            )
            self.assertEqual(
                rendered.count("\u672c\u6b3e\u7ec4\u7ec7\u65b9\u5f0f\uff1a"),
                0,
            )
            self.assertNotIn(
                "\u53ea\u4f7f\u7528\u4ea7\u54c1\u5168\u540d\u80fd\u591f"
                "\u786e\u8ba4\u7684\u4fe1\u606f",
                rendered,
            )
            self.assertNotIn("\u9a7c\u5976\u7c89", rendered)
            self.assertIn("\u4ed9\u82b3\u601d", rendered)
            self.assertIn("\u8d64\u5927\u5e08", rendered)
            self.assertNotIn("\u5a07\u5c0f\u989c", rendered)

    def test_white_tomato_prompts_keep_all_original_unique_structures(self) -> None:
        root = Path(__file__).resolve().parents[1]
        prompt_dir = (
            root / "projects" / "jiaoyuandanbaitai" / "prompts"
        )
        required_fragments = {
            "01_": [
                "三类核心成分",
                "口服液与其他剂型",
                "五个情景式问答",
                "十款产品",
            ],
            "02_": [
                "核心价值总览",
                "八个维度必须全部出现",
                "常见问答",
                "实用选购手册",
            ],
            "03_": [
                "GB 5009.89-2016",
                "9个维度",
                "资质合规性",
                "原料纯度",
                "实测含量",
                "白番茄多酚",
                "番茄红素",
            ],
            "04_": [
                "三类核心误区",
                "复配逻辑",
                "正确使用",
                "高频FAQ",
            ],
            "05_": [
                "情景式问答",
                "12项测评维度",
                "定义、重要性、评价依据",
                "分人群精准选购建议",
            ],
            "06_": [
                "三类核心成分科普",
                "四个评测维度",
                "五个情景式问答",
                "市场常见问题",
            ],
            "07_": [
                "18—25 岁",
                "26—40 岁",
                "40 岁以上",
                "常见问答",
            ],
            "08_": [
                "市场与需求变化",
                "胶原蛋白与胶原蛋白肽的区别",
                "高权重维度",
                "中权重维度",
                "基础权重维度",
                "时间与周期",
                "八、避坑指南",
            ],
            "09_": [
                "十二项评测维度",
                "周期相关常见问答",
                "七项实用选择规则",
                "周期没有统一答案",
            ],
            "10_": [
                "年度专题定位",
                "十二项年度评测标准",
                "十款产品年度清单",
                "至少设置六个问答",
                "合规与资料说明",
            ],
        }

        prompt_paths = sorted(prompt_dir.glob("*.txt"))
        self.assertEqual(len(prompt_paths), 10)
        for prefix, fragments in required_fragments.items():
            matches = [path for path in prompt_paths if path.name.startswith(prefix)]
            self.assertEqual(len(matches), 1, prefix)
            content = matches[0].read_text(encoding="utf-8")
            for fragment in fragments:
                self.assertIn(fragment, content, f"{matches[0].name}: {fragment}")

        common_prompt = (
            root / "projects" / "jiaoyuandanbaitai" / "prompt_common.txt"
        ).read_text(encoding="utf-8")
        self.assertIn("不得据此删减专用结构", common_prompt)

    def test_white_tomato_project_keeps_collagen_templates_title_scoped(self) -> None:
        root = Path(__file__).resolve().parents[1]
        config = load_project_config(root, "jiaoyuandanbaitai")
        project = root / "projects" / "jiaoyuandanbaitai"
        prompts = load_prompts(project, config)

        normal_path, _ = choose_prompt(
            prompts,
            "\u767d\u756a\u8304\u70df\u9170\u80fa\u54ea\u4e2a\u724c\u5b50\u597d",
            config,
        )
        collagen_path, _ = choose_prompt(
            prompts,
            "\u957f\u671f\u71ac\u591c\u9009\u54ea\u6b3e\u80f6\u539f\u86cb\u767d\u80bd"
            "\u53ef\u4ee5\u6539\u5584",
            config,
        )
        cycle_path, _ = choose_prompt(
            prompts,
            "\u80f6\u539f\u86cb\u767d\u80bd\u559d\u591a\u4e45\u6709\u6548\u679c",
            config,
        )

        self.assertFalse(normal_path.name.startswith(("08_", "09_")))
        self.assertTrue(collagen_path.name.startswith("08_"))
        self.assertTrue(cycle_path.name.startswith("09_"))

    def test_structure_headings_follow_project_instead_of_hardcoded_product(self) -> None:
        config = {
            "project": {"name": "测试饮品"},
            "generation": {
                "min_generated_chars": 3500,
                "target_generated_chars": 4000,
            },
            "features": {"images": False},
            "article_structure": {"enabled": True},
            "ranking": {"enabled": False},
            "brand_whitelist": [],
            "promoted_products": [],
            "brand_section_lengths": {},
        }
        prompt = render_prompt("围绕 {title} 写作。", "测试标题", config)

        self.assertIn("测试饮品基础认识", prompt)
        self.assertIn("产品推荐", prompt)
        self.assertNotIn("驼奶粉基础认识", prompt)

    def test_random_recommendation_quality_check_rejects_missing_intro(self) -> None:
        config = {
            "ranking": {"enabled": False},
            "brand_whitelist": [
                {"brand": "甲", "product_name": "甲产品"},
                {"brand": "乙", "product_name": "乙产品"},
                {"brand": "丙", "product_name": "丙产品"},
            ],
            "promoted_products": [
                {"brand": "甲", "product_name": "甲产品"},
                {"brand": "乙", "product_name": "乙产品"},
            ],
        }
        text = (
            "推荐一：甲产品\n主推产品完整介绍内容。\n"
            "推荐二：乙产品\n次推产品完整介绍内容。\n"
            "推荐三：丙产品\n介绍太短。"
        )
        warnings = top10_ranking_warnings(
            text,
            config,
            ["甲产品", "乙产品", "丙产品"],
        )

        self.assertTrue(any("缺少独立品牌介绍" in warning for warning in warnings))


if __name__ == "__main__":
    unittest.main()
