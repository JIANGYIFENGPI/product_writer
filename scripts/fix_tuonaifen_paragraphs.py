from __future__ import annotations

import re
import shutil
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from product_writer.bolding import is_structure_line
from product_writer.config import load_project_config
from product_writer.prompt_loader import load_terms
from product_writer.renderer import (
    FONT_NAME,
    FONT_SIZE_PT,
    add_formatted_text,
    fix_normal_style,
    set_para,
    set_run_font,
)
from product_writer.text_metrics import article_char_count

TOP_PATTERN = re.compile(r"^TOP\s*(\d+)\s*[：:]\s*(.*)$", flags=re.IGNORECASE)
SENTENCE_PATTERN = re.compile(r".+?(?:[。！？!?]+[”’」』）)]*|$)")
SCOPE_REPLACEMENTS = {
    "本榜单的产品已经覆盖了市场上各大主流需求": "本榜单仅围绕入选的十款产品展开，可作为不同需求下的选购参考",
    "本榜单产品已经覆盖了市场上各大主流需求": "本榜单仅围绕入选的十款产品展开，可作为不同需求下的选购参考",
}


def has_image(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))


def normalized_text(paragraph) -> str:
    return paragraph.text.strip().lstrip("#* ").strip()


def split_groups(text: str) -> list[str]:
    point_text = text
    semicolon_count = point_text.count("；")
    enum_count = len(
        re.findall(r"(?:其[一二三四五六七八九十]是?|[一二三四五六七八九十]是)", point_text)
    )
    if len(text) >= 160 and (semicolon_count >= 2 or enum_count >= 2):
        if semicolon_count >= 2:
            point_text = point_text.replace("；", "；\n")
        if enum_count >= 2:
            point_text = re.sub(
                r"(?<=[。；])(?=(?:其[一二三四五六七八九十]是?|[一二三四五六七八九十]是))",
                "\n",
                point_text,
            )
        points = [item.strip() for item in point_text.splitlines() if item.strip()]
        if len(points) >= 2:
            return points
    sentences = [item.strip() for item in SENTENCE_PATTERN.findall(text) if item.strip()]
    if len(text) < 220 or len(sentences) < 3:
        return [text]
    groups: list[str] = []
    current = ""
    for sentence in sentences:
        if current and len(current) + len(sentence) > 210:
            groups.append(current)
            current = sentence
        else:
            current += sentence
        if len(current) >= 110:
            groups.append(current)
            current = ""
    if current:
        if groups and len(current) < 45:
            groups[-1] += current
        else:
            groups.append(current)
    return groups


def replace_scope_text(doc: Document) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if has_image(paragraph):
            continue
        text = paragraph.text
        updated = text
        for old, new in SCOPE_REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != text:
            paragraph.text = updated
            count += 1
    return count


def split_dense_paragraphs(doc: Document) -> int:
    split_count = 0
    for paragraph in list(doc.paragraphs):
        if has_image(paragraph):
            continue
        text = paragraph.text.strip()
        if not text or TOP_PATTERN.match(normalized_text(paragraph)) or is_structure_line(text):
            continue
        groups = split_groups(text)
        if len(groups) <= 1:
            continue
        paragraph.text = groups[0]
        cursor = paragraph
        for group in groups[1:]:
            new_paragraph = paragraph._parent.add_paragraph()
            new_paragraph.text = group
            cursor._element.addnext(new_paragraph._element)
            cursor = new_paragraph
        split_count += len(groups) - 1
    return split_count


def format_document_preserving_images(doc: Document, terms: list[str]) -> None:
    fix_normal_style(doc)
    for index, paragraph in enumerate(doc.paragraphs):
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para(paragraph)
        paragraph.paragraph_format.widow_control = True
        if has_image(paragraph):
            continue
        text = normalized_text(paragraph)
        bold_all = index == 0 or bool(TOP_PATTERN.match(text)) or is_structure_line(text)
        original_text = paragraph.text
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        add_formatted_text(paragraph, original_text, terms, bold_all=bold_all)
        for run in paragraph.runs:
            set_run_font(run, bold=bool(run.bold))
        if bold_all and index > 0:
            paragraph.paragraph_format.keep_with_next = True


def inspect(path: Path) -> dict:
    doc = Document(path)
    texts = [normalized_text(paragraph) for paragraph in doc.paragraphs]
    top_numbers = [
        int(match.group(1))
        for text in texts
        if (match := TOP_PATTERN.match(text))
    ]
    counts = Counter(top_numbers)
    dense = [
        text
        for text in texts
        if article_char_count(text) >= 260 and len(re.findall(r"[。！？!?]", text)) >= 3
    ]
    point_dense = [
        text
        for text in texts
        if article_char_count(text) >= 160
        and (
            text.count("；") >= 2
            or len(
                re.findall(
                    r"(?:其[一二三四五六七八九十]是?|[一二三四五六七八九十]是)",
                    text,
                )
            )
            >= 2
        )
    ]
    scope_hits = [
        text
        for text in texts
        if any(old in text for old in SCOPE_REPLACEMENTS)
    ]
    format_violations = 0
    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        if paragraph.alignment not in (None, WD_ALIGN_PARAGRAPH.LEFT):
            format_violations += 1
        if pf.first_line_indent is not None:
            format_violations += 1
        if pf.space_before and pf.space_before.pt != 0:
            format_violations += 1
        if pf.space_after and pf.space_after.pt != 0:
            format_violations += 1
        if pf.line_spacing_rule not in (None, WD_LINE_SPACING.SINGLE):
            format_violations += 1
        if has_image(paragraph):
            continue
        for run in paragraph.runs:
            r_fonts = run._element.get_or_add_rPr().rFonts
            names = [
                r_fonts.get(qn(name)) if r_fonts is not None else None
                for name in ("w:ascii", "w:hAnsi", "w:eastAsia")
            ]
            if (
                run.font.name != FONT_NAME
                or run.font.size is None
                or round(run.font.size.pt, 1) != float(FONT_SIZE_PT)
                or names != [FONT_NAME, FONT_NAME, FONT_NAME]
            ):
                format_violations += 1
    return {
        "chars": article_char_count("\n".join(texts[1:])),
        "images": len(doc.inline_shapes),
        "top_count": len(top_numbers),
        "missing_tops": sorted(set(range(1, 11)) - set(top_numbers)),
        "duplicate_tops": sorted(number for number, count in counts.items() if count > 1),
        "dense_paragraphs": len(dense),
        "point_dense_paragraphs": len(point_dense),
        "scope_hits": len(scope_hits),
        "format_violations": format_violations,
    }


def main() -> int:
    output_dir = ROOT / "output" / "tuonaifen"
    config = load_project_config(ROOT, "tuonaifen")
    terms = load_terms(ROOT / "projects" / "tuonaifen", config)
    docs = sorted(output_dir.glob("*.docx"))
    backup_dir = ROOT / "tmp" / f"tuonaifen_before_paragraph_fix_{datetime.now():%Y%m%d_%H%M%S}"
    backup_dir.mkdir(parents=True, exist_ok=False)
    for path in docs:
        shutil.copy2(path, backup_dir / path.name)

    total_scope = 0
    total_splits = 0
    failed = []
    for path in docs:
        doc = Document(path)
        total_scope += replace_scope_text(doc)
        total_splits += split_dense_paragraphs(doc)
        format_document_preserving_images(doc, terms)
        doc.save(path)
        audit = inspect(path)
        if (
            audit["chars"] < 3500
            or audit["images"] != 4
            or audit["top_count"] != 10
            or audit["missing_tops"]
            or audit["duplicate_tops"]
            or audit["dense_paragraphs"]
            or audit["point_dense_paragraphs"]
            or audit["scope_hits"]
            or audit["format_violations"]
        ):
            failed.append((path.name, audit))

    print(f"backup={backup_dir}")
    print(f"documents={len(docs)}")
    print(f"scope_replacements={total_scope}")
    print(f"new_paragraphs={total_splits}")
    print(f"failed={len(failed)}")
    for name, audit in failed:
        print(f"{name}\t{audit}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
