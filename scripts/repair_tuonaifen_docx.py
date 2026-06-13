from __future__ import annotations

import re
import shutil
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from product_writer.bolding import is_structure_line
from product_writer.config import load_project_config
from product_writer.prompt_loader import load_terms
from product_writer.renderer import (
    FONT_NAME,
    FONT_SIZE_PT,
    add_formatted_text,
    fix_normal_style,
    set_para,
    set_run_font,
)
from product_writer.text_metrics import article_char_count


TOP_PATTERN = re.compile(r"^TOP\s*(\d+)\s*[：:]\s*(.*)$", flags=re.IGNORECASE)
QUICK_LIST_HINTS = ("简榜", "速览", "一览", "名单", "排名", "排行榜")
PROMPT_TITLE_PATTERN = re.compile(r"^\s*(?:文章)?标题\s*[：:]\s*(.+?)\s*$")


def normalized_paragraph_text(paragraph) -> str:
    return paragraph.text.strip().lstrip("#* ").strip()


def top_match(text: str):
    return TOP_PATTERN.match(text.strip())


def find_top_runs(texts: list[str]) -> list[list[int]]:
    runs: list[list[int]] = []
    current: list[int] = []
    previous_number: int | None = None
    for index, text in enumerate(texts):
        match = top_match(text)
        if not match:
            continue
        number = int(match.group(1))
        if current and previous_number is not None and number != previous_number + 1:
            runs.append(current)
            current = []
        current.append(index)
        previous_number = number
        if number == 10:
            runs.append(current)
            current = []
            previous_number = None
    if current:
        runs.append(current)
    return runs


def is_quick_list_run(texts: list[str], run: list[int]) -> bool:
    numbers = [int(top_match(texts[index]).group(1)) for index in run if top_match(texts[index])]
    if numbers != list(range(1, 11)):
        return False
    gaps = [run[index + 1] - run[index] for index in range(len(run) - 1)]
    if gaps and max(gaps) <= 2:
        return True
    start = max(0, run[0] - 3)
    context = "".join(texts[start:run[0]])
    return any(hint in context for hint in QUICK_LIST_HINTS)


def quick_list_entries(texts: list[str], run: list[int]) -> dict[int, str]:
    entries: dict[int, str] = {}
    for index in run:
        match = top_match(texts[index])
        if not match:
            continue
        name = re.split(r"\s*[—–-]{2,}\s*", match.group(2), maxsplit=1)[0].strip()
        if name:
            entries[int(match.group(1))] = name
    return entries


def find_later_detail_headings(
    texts: list[str],
    run: list[int],
    entries: dict[int, str],
) -> dict[int, int]:
    headings: dict[int, int] = {}
    for rank, name in entries.items():
        top_token = f"TOP{rank}"
        for index in range(run[-1] + 1, len(texts)):
            text = texts[index]
            if len(text) > 140 or name not in text:
                continue
            match = top_match(text)
            exact_rank = bool(match and int(match.group(1)) == rank)
            embedded_rank = top_token.lower() in text.lower()
            brand_heading = text.startswith(name) and is_structure_line(text)
            if exact_rank or embedded_rank or brand_heading:
                headings[rank] = index
                break
    return headings


def remove_indices(doc: Document, indices: set[int]) -> None:
    for index in sorted(indices, reverse=True):
        paragraph = doc.paragraphs[index]
        paragraph._element.getparent().remove(paragraph._element)


def remove_images(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        drawings = paragraph._element.xpath(".//w:drawing | .//w:pict")
        if not drawings:
            continue
        removed += len(drawings)
        if not paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)
            continue
        for drawing in drawings:
            drawing.getparent().remove(drawing)
    return removed


def purge_image_relationships(doc: Document) -> int:
    image_relationship_ids = [
        rel_id
        for rel_id, relationship in doc.part.rels.items()
        if relationship.reltype == RT.IMAGE
    ]
    for rel_id in image_relationship_ids:
        doc.part.drop_rel(rel_id)
    return len(image_relationship_ids)


def load_obsolete_prompt_titles(project_dir: Path) -> set[str]:
    titles: set[str] = set()
    for path in (project_dir / "prompts").glob("*.txt"):
        for line in path.read_text(encoding="utf-8-sig").splitlines():
            match = PROMPT_TITLE_PATTERN.match(line)
            if match:
                titles.add(match.group(1).strip())
    return titles


def remove_obsolete_prompt_titles(doc: Document, obsolete_titles: set[str]) -> int:
    remove: set[int] = set()
    paragraphs = doc.paragraphs
    for index in range(1, min(6, len(paragraphs))):
        text = normalized_paragraph_text(paragraphs[index])
        if text in obsolete_titles:
            remove.add(index)
    remove_indices(doc, remove)
    return len(remove)


def split_long_top_paragraphs(doc: Document) -> int:
    split_count = 0
    for paragraph in list(doc.paragraphs):
        text = normalized_paragraph_text(paragraph)
        match = top_match(text)
        if not match:
            continue
        remainder = match.group(2).strip()
        separator = re.search(r"[。！？!?]", remainder)
        if not separator:
            continue
        heading = f"TOP{int(match.group(1))}：{remainder[:separator.start()].strip()}"
        body = remainder[separator.start():].lstrip("。！？!? ").strip()
        if not body:
            continue
        paragraph.text = heading
        body_paragraph = paragraph._parent.add_paragraph()
        paragraph._element.addnext(body_paragraph._element)
        body_paragraph.add_run(body)
        split_count += 1
    return split_count


def format_document(doc: Document, terms: list[str]) -> None:
    fix_normal_style(doc)
    for index, paragraph in enumerate(doc.paragraphs):
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para(paragraph)
        text = normalized_paragraph_text(paragraph)
        bold_all = index == 0 or bool(top_match(text)) or is_structure_line(text)
        original_text = paragraph.text
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        add_formatted_text(paragraph, original_text, terms, bold_all=bold_all)
        for run in paragraph.runs:
            set_run_font(run, bold=bool(run.bold))


def repair_document(path: Path, terms: list[str], obsolete_titles: set[str]) -> dict:
    doc = Document(path)
    removed_images = remove_images(doc)
    removed_image_relationships = purge_image_relationships(doc)
    removed_obsolete_titles = remove_obsolete_prompt_titles(doc, obsolete_titles)
    split_count = split_long_top_paragraphs(doc)
    texts = [normalized_paragraph_text(paragraph) for paragraph in doc.paragraphs]
    runs = find_top_runs(texts)
    complete_runs = [run for run in runs if len(run) == 10]
    removed_quick_lines = 0
    if complete_runs:
        quick_runs: list[list[int]] = []
        detail_headings: dict[int, tuple[int, str]] = {}
        for run in complete_runs:
            if not is_quick_list_run(texts, run):
                continue
            entries = quick_list_entries(texts, run)
            later_headings = find_later_detail_headings(texts, run, entries)
            if len(later_headings) != 10:
                continue
            quick_runs.append(run)
            for rank, index in later_headings.items():
                detail_headings[rank] = (index, entries[rank])

        for rank, (index, name) in detail_headings.items():
            doc.paragraphs[index].text = f"TOP{rank}：{name}"
        remove = {index for run in quick_runs for index in run}
        for run in quick_runs:
            heading_index = run[0] - 1
            if heading_index >= 1 and any(hint in texts[heading_index] for hint in QUICK_LIST_HINTS):
                remove.add(heading_index)
        removed_quick_lines = len(remove)
        remove_indices(doc, remove)
    format_document(doc, terms)
    doc.save(path)
    return {
        "removed_images": removed_images,
        "removed_image_relationships": removed_image_relationships,
        "removed_obsolete_titles": removed_obsolete_titles,
        "removed_quick_lines": removed_quick_lines,
        "split_top_paragraphs": split_count,
    }


def inspect_document(path: Path, whitelist: set[str]) -> dict:
    doc = Document(path)
    texts = [normalized_paragraph_text(paragraph) for paragraph in doc.paragraphs]
    top_numbers = [
        int(match.group(1))
        for text in texts
        if (match := top_match(text))
    ]
    counts = Counter(top_numbers)
    duplicate_tops = sorted(number for number, count in counts.items() if count > 1)
    missing_tops = sorted(set(range(1, 11)) - set(top_numbers))
    unknown_top_lines = [
        text
        for text in texts
        if top_match(text) and not any(name in text for name in whitelist)
    ]
    format_violations = 0
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        if paragraph.alignment not in (None, WD_ALIGN_PARAGRAPH.LEFT):
            format_violations += 1
        if paragraph_format.first_line_indent is not None:
            format_violations += 1
        if paragraph_format.space_before and paragraph_format.space_before.pt != 0:
            format_violations += 1
        if paragraph_format.space_after and paragraph_format.space_after.pt != 0:
            format_violations += 1
        if paragraph_format.line_spacing_rule not in (None, WD_LINE_SPACING.SINGLE):
            format_violations += 1
        for run in paragraph.runs:
            r_fonts = run._element.get_or_add_rPr().rFonts
            names = [
                r_fonts.get(qn(name)) if r_fonts is not None else None
                for name in ("w:ascii", "w:hAnsi", "w:eastAsia")
            ]
            if (
                run.font.name != FONT_NAME
                or run.font.size is None
                or round(run.font.size.pt, 1) != float(FONT_SIZE_PT)
                or names != [FONT_NAME, FONT_NAME, FONT_NAME]
            ):
                format_violations += 1
    return {
        "chars": article_char_count("\n".join(texts[1:])),
        "images": len(doc.inline_shapes),
        "top_count": len(top_numbers),
        "missing_tops": missing_tops,
        "duplicate_tops": duplicate_tops,
        "unknown_top_lines": unknown_top_lines,
        "format_violations": format_violations,
    }


def main() -> int:
    root = ROOT
    output_dir = root / "output" / "tuonaifen"
    config = load_project_config(root, "tuonaifen")
    terms = load_terms(root / "projects" / "tuonaifen", config)
    obsolete_titles = load_obsolete_prompt_titles(root / "projects" / "tuonaifen")
    whitelist = {
        value
        for item in config.get("brand_whitelist") or []
        for value in (item.get("brand"), item.get("product_name"))
        if value
    }
    docs = sorted(output_dir.glob("*.docx"))
    backup_dir = root / "tmp" / f"tuonaifen_docx_backup_{datetime.now():%Y%m%d_%H%M%S}"
    backup_dir.mkdir(parents=True, exist_ok=False)
    for path in docs:
        shutil.copy2(path, backup_dir / path.name)

    repaired = []
    for path in docs:
        changes = repair_document(path, terms, obsolete_titles)
        audit = inspect_document(path, whitelist)
        repaired.append((path.name, changes, audit))

    print(f"backup={backup_dir}")
    print(f"documents={len(repaired)}")
    for name, changes, audit in repaired:
        print(
            f"{name}\timages_removed={changes['removed_images']}"
            f"\timage_relationships_removed={changes['removed_image_relationships']}"
            f"\tobsolete_titles_removed={changes['removed_obsolete_titles']}"
            f"\tquick_lines_removed={changes['removed_quick_lines']}"
            f"\ttop_splits={changes['split_top_paragraphs']}"
            f"\tchars={audit['chars']}\timages={audit['images']}"
            f"\ttop_count={audit['top_count']}"
            f"\tmissing={audit['missing_tops']}"
            f"\tduplicates={audit['duplicate_tops']}"
            f"\tunknown={len(audit['unknown_top_lines'])}"
            f"\tformat_violations={audit['format_violations']}"
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
