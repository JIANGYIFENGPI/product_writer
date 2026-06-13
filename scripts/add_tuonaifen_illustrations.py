from __future__ import annotations

import csv
import hashlib
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from product_writer.bolding import is_structure_line
from product_writer.config import load_project_config
from product_writer.renderer import set_para, set_run_font


TOP_PATTERN = re.compile(r"^TOP\s*(\d+)\s*[：:]", flags=re.IGNORECASE)
TOPIC_TERMS = {
    "低温": ("低温", "工艺", "process", "temperature"),
    "有机": ("有机", "牧场", "pasture", "source", "clean"),
    "正宗": ("牧场", "奶源", "pasture", "source", "camel"),
    "纯正": ("牧场", "奶源", "pasture", "source", "clean"),
    "高钙": ("钙", "calcium", "骨"),
    "蛋白": ("蛋白", "protein", "营养"),
    "吸收": ("吸收", "消化", "digestive", "absorption"),
    "肠胃": ("肠胃", "消化", "digestive", "absorption"),
    "儿童": ("儿童", "孩子", "child", "children", "study", "blocks"),
    "青少年": ("青少年", "学习", "study", "school", "learning"),
    "中老年": ("中老年", "老人", "senior"),
    "睡眠": ("睡眠", "夜", "night", "bedtime"),
    "全家": ("全家", "家庭", "family"),
    "送礼": ("送礼", "礼盒", "gift"),
    "无糖": ("低糖", "控糖", "low"),
    "安全": ("配料", "认证", "溯源", "clean", "source"),
    "营养": ("营养", "nutrition", "protein", "calcium"),
}
SCENE_TRANSLATIONS = {
    "reading still life": "阅读静物",
    "family breakfast": "家庭早餐场景",
    "window milk": "窗边冲调场景",
    "camping packshot": "户外产品图",
    "camel pasture panorama": "骆驼牧场全景",
    "camel source": "骆驼奶源",
    "pasture source": "牧场奶源",
    "packshot dolls": "产品展示图",
}
DEFAULT_LONG_PREAMBLE_CHARS = 300
DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES = 2


def load_catalog(
    project_dir: Path,
    slot_aliases: dict[str, str] | None = None,
) -> list[dict[str, str | Path]]:
    rows: list[dict[str, str | Path]] = []
    aliases = slot_aliases or {}
    with (project_dir / "image_catalog.csv").open(encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            slot = aliases.get(row["slot"], row["slot"])
            path = project_dir / "images" / slot / row["file"]
            if path.exists():
                rows.append({**row, "slot": slot, "path": path})
    return rows


def topic_tokens(title: str) -> tuple[str, ...]:
    tokens: list[str] = []
    for keyword, values in TOPIC_TERMS.items():
        if keyword in title:
            tokens.extend(values)
    if not tokens:
        tokens.extend(("营养", "nutrition", "牧场", "pasture", "产品"))
    return tuple(dict.fromkeys(token.lower() for token in tokens))


def image_score(row: dict[str, str | Path], tokens: tuple[str, ...], title: str) -> int:
    haystack = f"{row['file']} {row['scene']}".lower()
    score = sum(8 for token in tokens if token in haystack)
    if any(word in haystack for word in ("信息图", "卡片", "infographic", "feature", "banner")):
        score += 3
    digest = hashlib.sha256(f"{title}|{row['file']}".encode("utf-8")).digest()
    return score * 100000 + int.from_bytes(digest[:2], "big")


def choose_images(
    catalog: list[dict[str, str | Path]],
    title: str,
    slots: list[str],
) -> list[dict[str, str | Path]]:
    tokens = topic_tokens(title)
    chosen: list[dict[str, str | Path]] = []
    for slot in slots:
        choices = [row for row in catalog if row["slot"] == slot]
        choices.sort(key=lambda row: image_score(row, tokens, title), reverse=True)
        if choices:
            chosen.append(choices[0])
    return chosen


def compress_image(source: Path, cache_dir: Path) -> Path:
    cache_dir.mkdir(parents=True, exist_ok=True)
    target = cache_dir / f"{source.stem}.jpg"
    if target.exists():
        return target
    with Image.open(source) as image:
        image = image.convert("RGB")
        image.thumbnail((1600, 1100), Image.Resampling.LANCZOS)
        image.save(target, "JPEG", quality=88, optimize=True, progressive=True)
    return target


def top_positions(doc: Document) -> dict[int, int]:
    positions: dict[int, int] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        match = TOP_PATTERN.match(paragraph.text.strip())
        if match:
            positions[int(match.group(1))] = index
    return positions


def is_quick_top_heading(doc: Document, index: int) -> bool:
    paragraphs = list(doc.paragraphs)
    text = paragraphs[index].text.strip()
    match = TOP_PATTERN.match(text)
    if not match:
        return False
    previous_top = index > 0 and bool(TOP_PATTERN.match(paragraphs[index - 1].text.strip()))
    next_top = index + 1 < len(paragraphs) and bool(TOP_PATTERN.match(paragraphs[index + 1].text.strip()))
    return previous_top or next_top


def detail_heading_index(doc: Document, rank: int, aliases: set[str]) -> int | None:
    paragraphs = list(doc.paragraphs)
    candidates: list[int] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        match = TOP_PATTERN.match(text)
        exact_top = bool(match and int(match.group(1)) == rank)
        product_heading = any(alias in text for alias in aliases) and (
            is_structure_line(text) or len(text) <= 120
        )
        if not exact_top and not product_heading:
            continue
        if exact_top and is_quick_top_heading(doc, index):
            continue
        following = [
            paragraphs[next_index].text.strip()
            for next_index in range(index + 1, min(index + 4, len(paragraphs)))
        ]
        if any(len(next_text) >= 80 and not TOP_PATTERN.match(next_text) for next_text in following):
            candidates.append(index)
    if candidates:
        return candidates[0]

    # 有些文章先列完整简榜，后文详情标题只写“第X章 + 产品名”。
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if any(alias in text for alias in aliases) and len(text) <= 140:
            following = [
                paragraphs[next_index].text.strip()
                for next_index in range(index + 1, min(index + 4, len(paragraphs)))
            ]
            if any(len(next_text) >= 80 for next_text in following):
                return index
    return None


def neutral_anchor(doc: Document, first_detail_index: int, config: dict):
    paragraphs = list(doc.paragraphs)
    settings = config.get("illustrations") or {}
    long_preamble_chars = int(
        settings.get("neutral_after_first_paragraph_when_preamble_chars")
        or DEFAULT_LONG_PREAMBLE_CHARS
    )
    min_paragraphs = int(
        settings.get("minimum_body_paragraphs_between_images")
        or DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES
    )
    body_indices = []
    for index in range(1, first_detail_index):
        text = paragraphs[index].text.strip()
        if len(text) >= 80 and not is_structure_line(text):
            body_indices.append(index)
    preamble_chars = sum(len(paragraphs[index].text.strip()) for index in body_indices)

    if (
        body_indices
        and len(body_indices) >= min_paragraphs + 1
        and preamble_chars >= long_preamble_chars
    ):
        return paragraphs[body_indices[0]], "after"

    # 短前言时将通用图放在首段前，即标题之后，避免靠近 TOP1 产品图。
    if body_indices:
        return paragraphs[body_indices[0]], "before"
    return paragraphs[0], "after"


def insert_illustration(
    anchor,
    image_path: Path,
    alt_text: str,
    *,
    keep_anchor: bool = False,
    position: str = "after",
) -> None:
    document = anchor._parent
    if keep_anchor:
        anchor.paragraph_format.keep_with_next = True
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(image_paragraph)
    run = image_paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(5.6))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", alt_text)
    doc_pr.set("descr", alt_text)

    if position == "before":
        anchor._element.addprevious(image_paragraph._element)
    else:
        anchor._element.addnext(image_paragraph._element)


def body_paragraphs_between_images(doc: Document, first_image: int, second_image: int) -> int:
    paragraphs = list(doc.paragraphs)
    return sum(
        1
        for paragraph in paragraphs[first_image + 1:second_image]
        if len(paragraph.text.strip()) >= 60 and not is_structure_line(paragraph.text.strip())
    )


def validate_image_spacing(doc: Document, minimum: int) -> None:
    image_indices = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph._element.xpath(".//w:drawing | .//w:pict")
    ]
    if len(image_indices) != 4:
        raise RuntimeError(f"插图数量异常：{len(image_indices)} != 4")
    if body_paragraphs_between_images(doc, image_indices[0], image_indices[1]) < minimum:
        raise RuntimeError("通用图与TOP1产品图之间正文不足，停止输出以避免两图过近")


def main() -> int:
    project_dir = ROOT / "projects" / "tuonaifen"
    output_dir = ROOT / "output" / "tuonaifen"
    config = load_project_config(ROOT, "tuonaifen")
    promoted_aliases: dict[int, set[str]] = {}
    for item in config.get("promoted_products") or []:
        rank = int(item.get("rank") or 0)
        promoted_aliases[rank] = {
            str(value)
            for value in (
                item.get("brand"),
                item.get("product_name"),
                *(item.get("aliases") or []),
            )
            if value
        }
    promoted_slots = {
        int(item.get("rank") or 0): str(item.get("image_slot") or f"top{item.get('rank')}").strip()
        for item in config.get("promoted_products") or []
        if int(item.get("rank") or 0) in {1, 2, 3}
    }
    slot_aliases = {
        "top1": next(
            (
                str(item.get("image_slot"))
                for item in config.get("promoted_products") or []
                if str(item.get("brand") or "") == "赤大师" and item.get("image_slot")
            ),
            "top1",
        ),
        "top2": next(
            (
                str(item.get("image_slot"))
                for item in config.get("promoted_products") or []
                if str(item.get("brand") or "") == "迪奢思" and item.get("image_slot")
            ),
            "top2",
        ),
        "top3": next(
            (
                str(item.get("image_slot"))
                for item in config.get("promoted_products") or []
                if str(item.get("brand") or "") == "六驼" and item.get("image_slot")
            ),
            "top3",
        ),
    }
    catalog = load_catalog(project_dir, slot_aliases)
    docs = sorted(output_dir.glob("*.docx"))
    backup_dir = ROOT / "tmp" / f"tuonaifen_before_illustrations_{datetime.now():%Y%m%d_%H%M%S}"
    backup_dir.mkdir(parents=True, exist_ok=False)
    cache_dir = ROOT / "tmp" / "tuonaifen_illustration_cache"

    for path in docs:
        shutil.copy2(path, backup_dir / path.name)

    for path in docs:
        doc = Document(path)
        if doc.inline_shapes:
            raise RuntimeError(f"文档已有图片，停止避免重复插图：{path.name}")
        chosen = choose_images(
            catalog,
            path.stem,
            [
                str((config.get("illustrations") or {}).get("neutral_slot") or "neutral"),
                *(promoted_slots[rank] for rank in (1, 2, 3)),
            ],
        )
        detail_indices = {
            rank: detail_heading_index(doc, rank, promoted_aliases.get(rank, set()))
            for rank in (1, 2, 3)
        }
        if any(index is None for index in detail_indices.values()):
            raise RuntimeError(f"无法识别TOP1-TOP3真实详情标题：{path.name} {detail_indices}")
        neutral, neutral_position = neutral_anchor(
            doc,
            min(index for index in detail_indices.values() if index is not None),
            config,
        )
        anchors = [
            neutral,
            doc.paragraphs[detail_indices[1]],
            doc.paragraphs[detail_indices[2]],
            doc.paragraphs[detail_indices[3]],
        ]
        if any(anchor is None for anchor in anchors) or len(chosen) != 4:
            raise RuntimeError(f"无法为文档找到4个安全插图位置：{path.name}")
        for index, (anchor, row) in enumerate(zip(anchors, chosen)):
            image_path = compress_image(Path(row["path"]), cache_dir)
            scene = SCENE_TRANSLATIONS.get(str(row["scene"]).lower(), str(row["scene"]))
            alt_text = f"{row['brand']}，{scene}"
            insert_illustration(
                anchor,
                image_path,
                alt_text,
                keep_anchor=index > 0,
                position=neutral_position if index == 0 else "after",
            )
        minimum = int(
            (config.get("illustrations") or {}).get("minimum_body_paragraphs_between_images")
            or DEFAULT_MIN_PARAGRAPHS_BETWEEN_IMAGES
        )
        validate_image_spacing(doc, minimum)
        doc.save(path)
        print(f"{path.name}\timages=4\t" + " | ".join(str(row["scene"]) for row in chosen))

    print(f"backup={backup_dir}")
    print(f"documents={len(docs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
